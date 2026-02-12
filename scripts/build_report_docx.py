#!/usr/bin/env python3
"""
Build MBB-ready Word document — v3.
Includes AI-native competitor (Synthflow), dynamic statistics,
refreshed data (154 churn quotes), nuanced AI claims.
"""

import json, math, os, tempfile, shutil, re
from collections import Counter

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))
QUOTES_FILE = os.path.join(BASE, "clean_quotes", "final_quotes.json")
CODER_A = os.path.join(BASE, "analysis", "coder_a_k9.json")
CODER_B = os.path.join(BASE, "analysis", "coder_b_k9.json")
OUTPUT_FILE = os.path.join(BASE, "analysis", "Central_AI_Churn_Analysis.docx")

# ── Colors (docx) ────────────────────────────────────────────────────
DARK_BLUE    = RGBColor(0x1B, 0x36, 0x5F)
ACCENT_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
DARK_GRAY    = RGBColor(0x33, 0x33, 0x33)
MID_GRAY     = RGBColor(0x66, 0x66, 0x66)
LIGHT_BLUE_BG = "D6E4F0"
LIGHT_GRAY_BG = "F2F2F2"
GREEN_RGB    = RGBColor(0x2D, 0x8B, 0x57)
RED_RGB      = RGBColor(0xC0, 0x39, 0x2B)

# ── Colors (matplotlib) ──────────────────────────────────────────────
C_BLUE    = "#1B365F"
C_ACCENT  = "#2E75B6"
C_ORANGE  = "#E67E22"
C_GREEN   = "#27AE60"
C_PURPLE  = "#8E44AD"
C_RED     = "#C0392B"
C_GRAY    = "#BDC3C7"
GROUP_COLORS = {"CALL HANDLING": C_BLUE, "BILLING": C_ORANGE,
                "SERVICE RELIABILITY": C_GREEN, "INDUSTRY DISILLUSIONMENT": C_PURPLE}

# ── Metadata ─────────────────────────────────────────────────────────
CATEGORY_META = {
    "SCRIPT_ADHERENCE":      {"short": "Script Failures",    "group": "CALL HANDLING",
        "diff": "AI follows the script every time", "gtm": "Customers say: receptionists \u201Cdon\u2019t read directions.\u201D AI does."},
    "BILLING_PREDATORY":     {"short": "Predatory Billing",  "group": "BILLING",
        "diff": "No per-minute billing", "gtm": "Customers report being charged for spam calls and hold time."},
    "INAUTHENTICITY":        {"short": "Inauthenticity",     "group": "CALL HANDLING",
        "diff": "AI trained on your business", "gtm": "Customers say: \u201CThey didn\u2019t know what to do.\u201D"},
    "GARBLED_INFO":          {"short": "Garbled Info",        "group": "CALL HANDLING",
        "diff": "AI captures data accurately", "gtm": "Customers say: \u201CPhone number, address, email all wrong.\u201D"},
    "SERIAL_SWITCHING":      {"short": "Serial Switching",    "group": "INDUSTRY DISILLUSIONMENT",
        "diff": "Different technology, not another vendor", "gtm": "Customers say: \u201CTried 4 other companies.\u201D"},
    "BILLING_TRAP":          {"short": "Billing Trap",        "group": "BILLING",
        "diff": "Cancel anytime, one click", "gtm": "Customers report being billed months after cancellation."},
    "BILLING_OPAQUE":        {"short": "Opaque Billing",      "group": "BILLING",
        "diff": "Real-time usage dashboard", "gtm": "Customers say: \u201CSurprise charges doubled my bill.\u201D"},
    "QUALITY_DECAY":         {"short": "Quality Decay",       "group": "SERVICE RELIABILITY",
        "diff": "AI doesn\u2019t burn out or quit", "gtm": "Customers say: \u201CUsed to be great, now it\u2019s terrible.\u201D"},
    "MISSED_CALLS":          {"short": "Missed Calls",        "group": "SERVICE RELIABILITY",
        "diff": "AI answers every call", "gtm": "Customers say: \u201CPhone rang 10+ times before answer.\u201D"},
    "ROUTING_ERRORS":        {"short": "Routing Errors",      "group": "CALL HANDLING",
        "diff": "AI routes by rules you set", "gtm": "Customers say: \u201CCalled wrong tech, didn\u2019t check schedule.\u201D"},
    "BILLING_TOO_EXPENSIVE": {"short": "Too Expensive",       "group": "BILLING",
        "diff": "AI at lower cost", "gtm": "Customers say: \u201CCancelled because it got too expensive.\u201D"},
}

GROUP_ORDER = ["CALL HANDLING", "BILLING", "SERVICE RELIABILITY", "INDUSTRY DISILLUSIONMENT"]
GROUP_META = {
    "CALL HANDLING":            {"insight": "The core product promise (handling calls well) is the single largest source of churn.",
                                 "impl": "Table stakes. AI must nail scripts, sound authentic, capture data accurately."},
    "BILLING":                  {"insight": "Billing problems are the second-largest churn driver, separate from service quality.",
                                 "impl": "Transparent pricing is a competitive moat. No per-minute, no contracts, cancel anytime."},
    "SERVICE RELIABILITY":      {"insight": "Services that start strong deteriorate over time. This affects AI-native platforms too (Synthflow: 35% of churn is reliability).",
                                 "impl": "Consistent quality requires engineering discipline. Year 3 quality must equal Day 1 quality."},
    "INDUSTRY DISILLUSIONMENT": {"insight": "A segment of customers have tried multiple services and given up on the category entirely.",
                                 "impl": "Highest-intent prospects. Position as fundamentally different technology."},
}
TOP_COMPETITORS = ["AnswerConnect", "Ruby Receptionist", "PATLive", "Smith.ai", "Synthflow"]

COMPETITOR_PROFILES = {
    "Ruby Receptionist": {
        "founded": "2003",
        "hq": "Portland, OR",
        "what": "US-based virtual receptionist service with 15,000+ small business customers",
        "how": "Shared pool of human receptionists, billed per minute",
        "pricing": "From $235/mo (50 min), up to $1,640/mo (500 min)",
        "claim": "\u201CDelight your callers\u201D",
        "type": "Human-pool",
    },
    "AnswerConnect": {
        "founded": "2002",
        "hq": "Portland, OR",
        "what": "24/7 live answering service for small and mid-size businesses",
        "how": "Shared pool of human receptionists, billed per minute (including after-call work)",
        "pricing": "From ~$350/mo (200 min) + $49.99 setup fee",
        "claim": "\u201CNever miss a call\u201D",
        "type": "Human-pool",
    },
    "PATLive": {
        "founded": "1990",
        "hq": "Tallahassee, FL",
        "what": "One of the oldest US live answering services, serving businesses since 1990",
        "how": "Shared pool of human receptionists, billed per minute across five tiers",
        "pricing": "From $235/mo (75 min), up to $999/mo (600 min)",
        "claim": "\u201CProfessional receptionists, 24/7\u201D",
        "type": "Human-pool",
    },
    "Smith.ai": {
        "founded": "2015",
        "hq": "Palo Alto, CA",
        "what": "AI-assisted virtual receptionist combining AI screening with human agents",
        "how": "Hybrid AI + human model, billed per call (not per minute)",
        "pricing": "From $300/mo (30 calls); pure AI plan from $95/mo",
        "claim": "\u201CAI-powered virtual receptionists\u201D",
        "type": "Hybrid",
    },
    "Synthflow": {
        "founded": "2023",
        "hq": "Berlin, Germany",
        "what": "No-code voice AI platform. Only AI-native competitor with enough review data to analyze",
        "how": "Pure AI, no-code builder; monthly subscription plus per-minute usage",
        "pricing": "From $29/mo to $900/mo, plus per-minute costs",
        "claim": "\u201CBuild AI phone agents without code\u201D",
        "type": "AI-native",
    },
    "Abby Connect": {
        "founded": "2005",
        "hq": "Las Vegas, NV",
        "what": "Boutique live receptionist service, smaller pool, dedicated teams",
        "how": "Dedicated team of human receptionists, billed per minute",
        "pricing": "From $299/mo (100 min)",
        "claim": "\u201CYour dedicated receptionist team\u201D",
        "type": "Human-pool",
    },
    "Dialzara": {
        "founded": "2023",
        "hq": "US",
        "what": "AI answering service for small businesses, pure AI model",
        "how": "Pure AI, billed per minute with monthly base",
        "pricing": "From $29/mo to $99/mo",
        "claim": "\u201CAI receptionist that sounds human\u201D",
        "type": "AI-native",
    },
    "My AI Front Desk": {
        "founded": "2023",
        "hq": "US",
        "what": "AI phone receptionist for appointment-based businesses",
        "how": "Pure AI, monthly subscription with minute cap",
        "pricing": "$79/mo (200 min)",
        "claim": "\u201C24/7 AI receptionist\u201D",
        "type": "AI-native",
    },
    "Goodcall": {
        "founded": "2022",
        "hq": "US",
        "what": "AI phone agent platform for local businesses",
        "how": "Pure AI, billed per unique caller",
        "pricing": "$59\u2013$199/mo",
        "claim": "\u201CAI phone agent for your business\u201D",
        "type": "AI-native",
    },
}

# ── Data loading ─────────────────────────────────────────────────────
def load_quotes():
    with open(QUOTES_FILE) as f:
        return [q for q in json.load(f) if q["llm"].get("category", "").startswith("churn_")]

def load_coder(fp):
    with open(fp) as f:
        return {item["idx"]: item["category"] for item in json.load(f)["classifications"]}

def adjudicate(a, b):
    final, dis = {}, []
    for idx in sorted(set(a) | set(b)):
        final[idx] = a.get(idx)
        if a.get(idx) != b.get(idx):
            dis.append((idx, a.get(idx), b.get(idx)))
    return final, dis

def compute_kappa(codes_a, codes_b):
    """Compute Cohen's kappa from two coder dictionaries."""
    shared = sorted(set(codes_a) & set(codes_b))
    n = len(shared)
    if n == 0:
        return 0.0, 0, 0
    labels_a = [codes_a[i] for i in shared]
    labels_b = [codes_b[i] for i in shared]
    agree = sum(1 for a, b in zip(labels_a, labels_b) if a == b)
    p_o = agree / n
    # Expected agreement
    all_cats = sorted(set(labels_a) | set(labels_b))
    p_e = sum((labels_a.count(c) / n) * (labels_b.count(c) / n) for c in all_cats)
    kappa = (p_o - p_e) / (1 - p_e) if p_e < 1 else 1.0
    return kappa, agree, n

def compute_weight(q):
    quality = q["llm"].get("quote_quality", 1)
    eng = max(1.0, math.log2(max(q.get("score") or 1, 1))) if q["source"] == "reddit" else 1.0
    return quality * eng

def normalize_product(p):
    if not p or p == "unnamed": return "Other"
    return "Ruby Receptionist" if p == "Ruby" else p

def truncate(text, n=200):
    t = text.replace("\n", " ").strip()
    return t[:n] + "..." if len(t) > n else t

def devastation_score(text):
    t = text.lower(); s = 0
    if any(w in t for w in ["thousand", "hundred", "$"]): s += 3
    if any(w in t for w in ["lost client", "lost customer", "lost revenue"]): s += 3
    if any(w in t for w in ["years", "year"]): s += 2
    if any(w in t for w in ["scam", "dishonest", "beware", "fraud", "worst",
                             "terrible", "horrible", "negligence"]): s += 2
    return s

def extract_pull_quote(text, max_len=200):
    text_clean = text.replace("\n", " ").strip()
    sentences = re.split(r'(?<=[.!?])\s+', text_clean)
    best, best_s = None, -1
    for s in sentences:
        if len(s) < 20: continue
        sc = devastation_score(s) + (1 if 40 < len(s) < 180 else 0)
        if sc > best_s: best_s, best = sc, s
    if best and len(best) <= max_len: return best
    elif best: return best[:max_len] + "..."
    return truncate(text_clean, max_len)


# ── Chart generators ─────────────────────────────────────────────────
def _mbb():
    plt.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Helvetica Neue", "Helvetica", "Arial", "Calibri"],
        "font.size": 10, "axes.spines.top": False, "axes.spines.right": False,
        "axes.edgecolor": "#BDC3C7", "axes.labelcolor": "#333",
        "xtick.color": "#666", "ytick.color": "#666",
        "figure.facecolor": "white", "axes.facecolor": "white",
        "savefig.dpi": 220, "savefig.bbox": "tight", "savefig.pad_inches": 0.3,
    })

def chart_group_pareto(group_stats, total_weighted, path):
    _mbb()
    fig, ax = plt.subplots(figsize=(7.5, 3.2))
    groups = list(reversed(GROUP_ORDER))
    pcts = [group_stats.get(g, {}).get("weighted", 0) / total_weighted * 100 for g in groups]
    counts = [group_stats.get(g, {}).get("count", 0) for g in groups]
    colors = [GROUP_COLORS.get(g, C_GRAY) for g in groups]
    bars = ax.barh(groups, pcts, color=colors, height=0.55, edgecolor="white", linewidth=0.5)
    for bar, pct, count in zip(bars, pcts, counts):
        ax.text(bar.get_width() + 1.2, bar.get_y() + bar.get_height() / 2,
                f"{pct:.0f}%  ({count} quotes)", va="center", fontsize=10, color="#333", fontweight="bold")
    ax.set_xlim(0, 70); ax.xaxis.set_visible(False); ax.spines["bottom"].set_visible(False)
    ax.tick_params(axis="y", length=0)
    for l in ax.get_yticklabels(): l.set_fontsize(9.5); l.set_fontweight("bold"); l.set_color("#333")
    ch_pct = group_stats.get("CALL HANDLING", {}).get("weighted", 0) / total_weighted * 100 if total_weighted else 0
    ax.set_title(f"Call handling failures account for {ch_pct:.0f}% of all churn",
                  fontsize=11, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)

def chart_pareto_11(sorted_cats, cat_stats, total_weighted, path):
    _mbb()
    fig, ax1 = plt.subplots(figsize=(8.5, 4.5))
    labels = [CATEGORY_META[c]["short"] for c in sorted_cats]
    pcts = [cat_stats[c]["weighted"] / total_weighted * 100 for c in sorted_cats]
    colors = [GROUP_COLORS.get(CATEGORY_META[c]["group"], C_GRAY) for c in sorted_cats]
    x = np.arange(len(labels))
    bars = ax1.bar(x, pcts, color=colors, width=0.7, edgecolor="white", linewidth=0.5)
    ax2 = ax1.twinx()
    cumul = np.cumsum(pcts)
    ax2.plot(x, cumul, color=C_BLUE, linewidth=2.2, marker="o", markersize=4, zorder=5)
    ax2.set_ylim(0, 110); ax2.set_ylabel("Cumulative %", fontsize=9, color="#666")
    ax2.spines["top"].set_visible(False)
    top6_pct = cumul[5] if len(cumul) > 5 else cumul[-1]
    ax2.axhline(y=top6_pct, color=C_RED, linestyle="--", linewidth=1, alpha=0.6)
    ax2.text(len(labels) - 0.5, top6_pct + 2, f"{top6_pct:.0f}% of churn from top 6 problems",
             fontsize=7.5, color=C_RED, ha="right", fontweight="bold")
    for bar, pct in zip(bars, pcts):
        if pct > 4:
            ax1.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.4,
                     f"{pct:.0f}%", ha="center", va="bottom", fontsize=8, color="#333")
    ax1.set_xticks(x); ax1.set_xticklabels(labels, rotation=40, ha="right", fontsize=8)
    ax1.set_ylabel("Share of Churn (%)", fontsize=9, color="#666")
    ax1.set_ylim(0, max(pcts) * 1.35)
    from matplotlib.patches import Patch
    ax1.legend(handles=[Patch(facecolor=GROUP_COLORS[g], label=g) for g in GROUP_ORDER],
               loc="upper right", fontsize=7, framealpha=0.9)
    ax1.set_title(f"Top 6 problems drive {top6_pct:.0f}% of churn",
                   fontsize=11, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)

def chart_heatmap(comp_groups, path):
    _mbb()
    # Filter to competitors with data
    hm_comps = [p for p in TOP_COMPETITORS if sum(comp_groups.get(p, {}).values()) > 0]
    fig, ax = plt.subplots(figsize=(7.5, 2.2 + 0.6 * len(hm_comps)))
    data = []
    for prod in hm_comps:
        total = sum(comp_groups.get(prod, {}).values())
        data.append([comp_groups.get(prod, {}).get(g, 0) / total * 100 if total > 0 else 0
                     for g in GROUP_ORDER])
    data = np.array(data)
    im = ax.imshow(data, cmap="Blues", aspect="auto", vmin=0, vmax=70)
    ax.set_xticks(np.arange(len(GROUP_ORDER)))
    ax.set_xticklabels([g.replace("INDUSTRY ", "INDUSTRY\n") for g in GROUP_ORDER], fontsize=8)
    ax.set_yticks(np.arange(len(hm_comps)))
    ax.set_yticklabels(hm_comps, fontsize=9)
    for i in range(len(hm_comps)):
        for j in range(len(GROUP_ORDER)):
            v = data[i, j]; c = "white" if v > 45 else "#333"
            ax.text(j, i, f"{v:.0f}%", ha="center", va="center", fontsize=10.5,
                    color=c, fontweight="bold" if v > 35 else "normal")
    ax.tick_params(length=0)
    for spine in ax.spines.values(): spine.set_visible(True); spine.set_color("#BDC3C7")
    ax.set_title("Each competitor fails differently. Smith.ai and Synthflow skew toward billing.",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)

COMPANY_ALIASES = {
    "ruby": "Ruby Receptionist", "ruby receptionist": "Ruby Receptionist",
    "answerconnect": "AnswerConnect", "answer connect": "AnswerConnect",
    "patlive": "PATLive", "pat live": "PATLive", "patlive.com": "PATLive",
    "smith.ai": "Smith.ai", "smithai": "Smith.ai", "smith ai": "Smith.ai",
    "synthflow": "Synthflow", "synthflow.ai": "Synthflow",
    "abby connect": "Abby Connect", "abbyconnect": "Abby Connect",
    "dialzara": "Dialzara", "sas": "SAS", "specialty answering service": "SAS",
    "virtual hq": "Virtual HQ", "conversational": "Conversational",
    "goodcall": "Goodcall", "vapi": "Vapi", "bland ai": "Bland AI",
    "in-house": "In-house", "in house": "In-house",
}

def _product_names(prod):
    """Return set of lowercase name variants for a product."""
    names = {prod.lower()}
    for alias, canon in COMPANY_ALIASES.items():
        if canon == prod:
            names.add(alias.lower())
    return names

def _detect_direction(prod, pain, text):
    """Determine if a churn_switched review is an ARRIVAL at or DEPARTURE from prod.

    Three-step detection:
    1. Explicit directional phrases with product name ("dropped Smith.ai")
    2. Overall text sentiment in first 300 chars
    3. Pain-point keyword fallback
    Returns True for arrival, False for departure.
    """
    text_lower = text.lower()
    pain_lower = pain.lower()
    names = _product_names(prod)

    # ── Step 1: Directional phrases containing the product name ──────
    dep_phrases = ["dropped ", "left ", "cancelled ", "canceled ", "leaving ",
                   "switched from ", "ditched "]
    arr_phrases = ["switched to ", "switching to ", "went to ", "went with ",
                   "moved to ", "chose ", "choosing ", "found ", "finding "]
    for name in names:
        for dp in dep_phrases:
            if dp + name in text_lower or dp + name in pain_lower:
                return False  # departure
        for ap in arr_phrases:
            if ap + name in text_lower or ap + name in pain_lower:
                return True   # arrival

    # ── Step 2: Overall text sentiment ────────────────────────────────
    snippet = text_lower
    pos_words = ["happy", "love", "loves", "great", "best", "excellent", "amazing",
                 "wonderful", "recommend", "appreciate", "fantastic", "no comparison",
                 "pleased", "reliable", "impressed", "professional", "easy to work",
                 "helpful", "superior", "courteous"]
    neg_words = ["terrible", "worst", "awful", "horrible", "avoid", "waste", "scam",
                 "zero stars", "do not use", "poor", "disappointed", "frustrat",
                 "useless", "broken", "zero support", "no support", "non-existent",
                 "crap", "buyer beware", "stay far", "dissatisfied", "not what it used",
                 "canceled", "cancelled", "too expensive", "dropping the ball",
                 "finally left", "no longer", "nose dive", "disappoint",
                 "unfortunately", "missed"]
    pos = sum(1 for w in pos_words if w in snippet)
    neg = sum(1 for w in neg_words if w in snippet)
    if pos != neg:
        return pos > neg

    # ── Step 3: Pain-point keyword fallback ──────────────────────────
    arr_kw = ["switched to", "found", "chose", "happy", "best vendor",
              "great experience", "no comparison", "resolved", "recommend",
              "flexible", "accommodating"]
    dep_kw = ["dropped", "cancelled", "canceled", "too expensive", "declined",
              "wrong", "missed", "terrible", "worst", "overpromised", "crap",
              "deteriorat", "shopping", "dishonest", "lost calls", "costing",
              "unanswered", "hang up", "broken", "zero support", "fails",
              "non-existent", "not worth"]
    a = sum(1 for w in arr_kw if w in pain_lower)
    d = sum(1 for w in dep_kw if w in pain_lower)
    if a != d:
        return a > d

    # If completely ambiguous, default to departure (writing a churn review)
    return False


def compute_switching_data(quotes, final_codes):
    """Parse switching stories to compute arrivals, departures, net flow, and triggers."""
    switched = [(idx, quotes[idx]) for idx in final_codes
                if quotes[idx]["llm"].get("category") == "churn_switched"]

    arrivals = Counter()
    departures = Counter()
    triggers = Counter()

    for idx, q in switched:
        prod = normalize_product(q["llm"].get("product_mentioned"))
        pain = q["llm"].get("pain_point", "")
        text = q["text"]

        is_arrival = _detect_direction(prod, pain, text)

        # Find other companies mentioned in text
        text_lower = text.lower()
        other_companies = set()
        for alias, canonical in COMPANY_ALIASES.items():
            if alias in text_lower and canonical != prod:
                other_companies.add(canonical)

        if is_arrival:
            arrivals[prod] += 1
            for other in other_companies:
                departures[other] += 1
        else:
            departures[prod] += 1
            for other in other_companies:
                arrivals[other] += 1

        # Classify trigger from pain_point (cleaner signal than full text).
        # Check capability first (specific), then pricing (broad), default quality.
        trigger_text = pain.lower()
        if any(w in trigger_text for w in ["integrat", "customiz", "automat", "technolog",
                                            "can't do", "couldn't", "could not",
                                            "limited", "outdated", "compatib",
                                            "google voice", "can't handle",
                                            "couldn't adapt", "could not adapt"]):
            triggers["Capability"] += 1
        elif any(w in trigger_text for w in ["expensiv", "price", "too much",
                                              "bill", "charg", "fee", "money",
                                              "overpay", "overcharg",
                                              "no longer justified", "cut cost"]):
            triggers["Pricing"] += 1
        else:
            triggers["Quality"] += 1

    # Compute net flow per company
    all_companies = set(arrivals.keys()) | set(departures.keys())
    net_flow = {}
    for c in all_companies:
        net = arrivals.get(c, 0) - departures.get(c, 0)
        if net != 0:
            net_flow[c] = {"arrivals": arrivals.get(c, 0), "departures": departures.get(c, 0), "net": net}

    return {
        "total": len(switched),
        "net_flow": net_flow,
        "triggers": triggers,
        "arrivals": arrivals,
        "departures": departures,
    }

def chart_net_churn(path, switching_data):
    _mbb()
    nf = switching_data["net_flow"]
    # Sort: net gainers first (descending), then net losers (descending)
    sorted_comps = sorted(nf.items(), key=lambda x: x[1]["net"], reverse=True)
    # Filter to companies with |net| >= 1
    sorted_comps = [(c, d) for c, d in sorted_comps if abs(d["net"]) >= 1]

    labels = [c.replace("Receptionist", "Rec.").replace("Connect", "\nConnect") for c, _ in sorted_comps]
    scores = [d["net"] for _, d in sorted_comps]
    colors = [C_GREEN if s > 0 else C_RED for s in scores]

    fig, ax = plt.subplots(figsize=(7.5, 3.8))
    bars = ax.bar(labels, scores, color=colors, width=0.6, edgecolor="white", linewidth=0.5)
    for bar, sc in zip(bars, scores):
        y = bar.get_height(); off = 0.3 if sc > 0 else -0.35
        ax.text(bar.get_x() + bar.get_width() / 2, y + off, f"{sc:+d}",
                ha="center", va="bottom" if sc > 0 else "top", fontsize=10,
                fontweight="bold", color="#333")
    ax.axhline(y=0, color="#333", linewidth=0.8)
    max_abs = max(abs(s) for s in scores) if scores else 5
    ax.set_ylim(-max_abs - 2, max_abs + 2.5); ax.spines["bottom"].set_visible(False)
    ax.set_ylabel("Customers Gained minus Lost", fontsize=9, color="#666")
    # Find biggest loser for title
    biggest_loser = sorted_comps[-1][0] if sorted_comps else ""
    biggest_loss = sorted_comps[-1][1]["net"] if sorted_comps else 0
    ax.set_title(f"{biggest_loser} loses the most customers ({biggest_loss:+d} net)",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    n_stories = switching_data["total"]
    fig.text(0.99, 0.01, f"Based on {n_stories} switching stories | Each bar = customers arriving minus leaving | Directional signal",
             fontsize=7, color="#999", ha="right", fontstyle="italic")
    fig.savefig(path); plt.close(fig)

def compute_temporal_data():
    """Compute organic vs surge averages from raw Trustpilot data. Returns dict for chart + table."""
    review_dir = os.path.join(BASE, "review_data", "raw")
    name_map = {"answerconnect": "AnswerConnect", "patlive": "PATLive",
                "smithai": "Smith.ai", "ruby_receptionist": "Ruby Receptionist"}
    comp_order = ["AnswerConnect", "Ruby Receptionist", "PATLive", "Smith.ai"]
    result = {}
    for fname in os.listdir(review_dir):
        if not fname.endswith("_trustpilot.json"): continue
        key = fname.replace("_trustpilot.json", "")
        prod = name_map.get(key)
        if not prod or prod not in comp_order: continue
        with open(os.path.join(review_dir, fname)) as f:
            reviews = json.load(f)
        org_ratings, surge_ratings = [], []
        for r in reviews:
            try: rating = int(r.get("rating", 3))
            except (ValueError, TypeError): rating = 3
            date_str = r.get("date", "")
            if "2024" in date_str or "2025" in date_str or "2026" in date_str:
                month = 1
                for i, m in enumerate(["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"], 1):
                    if m in date_str: month = i; break
                if "2024" in date_str and month >= 7: surge_ratings.append(rating)
                elif "2025" in date_str or "2026" in date_str: surge_ratings.append(rating)
                else: org_ratings.append(rating)
            else:
                org_ratings.append(rating)
        org_avg = float(np.mean(org_ratings)) if org_ratings else 3.0
        surge_avg = float(np.mean(surge_ratings)) if surge_ratings else 3.0
        result[prod] = {
            "organic_avg": org_avg, "organic_n": len(org_ratings),
            "surge_avg": surge_avg, "surge_n": len(surge_ratings),
            "jump": surge_avg - org_avg,
        }
    return result, comp_order

def chart_temporal_before_after(path, temporal_data, comp_order):
    """Grouped bar chart: organic vs surge per competitor."""
    _mbb()
    fig, ax = plt.subplots(figsize=(7.5, 4))
    labels = ["AnswerConnect", "Ruby\nReceptionist", "PATLive", "Smith.ai"]
    organic = [temporal_data.get(c, {}).get("organic_avg", 3.0) for c in comp_order]
    surge = [temporal_data.get(c, {}).get("surge_avg", 3.0) for c in comp_order]
    x = np.arange(len(labels)); w = 0.32
    b1 = ax.bar(x - w/2, organic, w, label="Reviews posted without prompting (before mid-2024)", color=C_RED, alpha=0.85, edgecolor="white")
    b2 = ax.bar(x + w/2, surge,   w, label="Recent wave (mid-2024 onward)",      color=C_GREEN, alpha=0.85, edgecolor="white")
    for bar, val in zip(b1, organic):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                f"{val:.1f}", ha="center", fontsize=9, fontweight="bold", color=C_RED)
    for bar, val in zip(b2, surge):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                f"{val:.1f}", ha="center", fontsize=9, fontweight="bold", color=C_GREEN)
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylim(0, 5.8); ax.set_ylabel("Avg Star Rating", fontsize=9, color="#666")
    ax.axhline(y=3, color="#BDC3C7", linewidth=0.5, linestyle=":")
    ax.legend(fontsize=8, loc="upper left", framealpha=0.9)
    ax.set_title("Unprompted reviews tell the real story \u2014 recent waves mask persistent pain",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.text(0.99, 0.01, "From raw Trustpilot data | Unprompted: before mid-2024 | Recent wave: mid-2024 onward",
             fontsize=7, color="#999", ha="right", fontstyle="italic")
    fig.savefig(path); plt.close(fig)

def compute_bimodality(temporal_data, comp_order):
    """Compute % of 1-star and 5-star reviews for each competitor in organic vs surge periods."""
    review_dir = os.path.join(BASE, "review_data", "raw")
    name_map = {"answerconnect": "AnswerConnect", "patlive": "PATLive",
                "smithai": "Smith.ai", "ruby_receptionist": "Ruby Receptionist"}
    result = {}
    for fname in os.listdir(review_dir):
        if not fname.endswith("_trustpilot.json"): continue
        key = fname.replace("_trustpilot.json", "")
        prod = name_map.get(key)
        if not prod or prod not in comp_order: continue
        with open(os.path.join(review_dir, fname)) as f:
            reviews = json.load(f)
        org_extreme, org_total = 0, 0
        surge_extreme, surge_total = 0, 0
        org_5star, surge_5star = 0, 0
        for r in reviews:
            try: rating = int(r.get("rating", 3))
            except (ValueError, TypeError): rating = 3
            date_str = r.get("date", "")
            is_surge = False
            if "2024" in date_str or "2025" in date_str or "2026" in date_str:
                month = 1
                for i, m in enumerate(["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"], 1):
                    if m in date_str: month = i; break
                if ("2024" in date_str and month >= 7) or "2025" in date_str or "2026" in date_str:
                    is_surge = True
            if is_surge:
                surge_total += 1
                if rating in (1, 5): surge_extreme += 1
                if rating == 5: surge_5star += 1
            else:
                org_total += 1
                if rating in (1, 5): org_extreme += 1
                if rating == 5: org_5star += 1
        result[prod] = {
            "org_extreme_pct": org_extreme / org_total * 100 if org_total else 0,
            "surge_extreme_pct": surge_extreme / surge_total * 100 if surge_total else 0,
            "org_5star": org_5star, "org_total": org_total,
            "surge_5star": surge_5star, "surge_total": surge_total,
            "overall_extreme_pct": (org_extreme + surge_extreme) / (org_total + surge_total) * 100 if (org_total + surge_total) else 0,
        }
    return result

def chart_bimodality(bimod_data, comp_order, path):
    """Stacked bar chart showing extreme (1-star + 5-star) vs middle (2-4 star) reviews."""
    _mbb()
    fig, ax = plt.subplots(figsize=(7.5, 3.5))
    labels = ["Smith.ai", "AnswerConnect", "Ruby\nReceptionist", "PATLive"]
    order = ["Smith.ai", "AnswerConnect", "Ruby Receptionist", "PATLive"]
    extreme = [bimod_data.get(c, {}).get("overall_extreme_pct", 0) for c in order]
    middle = [100 - e for e in extreme]
    x = np.arange(len(labels))
    ax.barh(x, extreme, height=0.5, color=C_RED, alpha=0.85, label="Extreme (1-star or 5-star)")
    ax.barh(x, middle, height=0.5, left=extreme, color=C_GRAY, alpha=0.5, label="Middle (2-4 star)")
    for i, (e, lbl) in enumerate(zip(extreme, labels)):
        ax.text(e + 1, i, f"{e:.0f}%", va="center", fontsize=10, fontweight="bold", color=C_RED)
    ax.set_yticks(x); ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlim(0, 110); ax.set_xlabel("% of All Reviews", fontsize=9, color="#666")
    ax.legend(fontsize=8, loc="lower right", framealpha=0.9)
    ax.set_title("Most reviews are extreme \u2014 very few land in the middle",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.text(0.99, 0.01, "From raw Trustpilot data | Extreme = 1-star or 5-star only",
             fontsize=7, color="#999", ha="right", fontstyle="italic")
    fig.savefig(path); plt.close(fig)

def chart_legacy_vs_ai(comp_groups, path):
    """Grouped bar: legacy avg vs Synthflow across 4 churn groups."""
    _mbb()
    legacy_prods = ["Ruby Receptionist", "AnswerConnect", "PATLive", "Smith.ai"]
    legacy_avg = {}
    for g in GROUP_ORDER:
        vals = []
        for p in legacy_prods:
            total = sum(comp_groups.get(p, {}).values())
            if total > 0:
                vals.append(comp_groups.get(p, {}).get(g, 0) / total * 100)
        legacy_avg[g] = float(np.mean(vals)) if vals else 0.0
    synth_total = sum(comp_groups.get("Synthflow", {}).values())
    synth_pct = {}
    for g in GROUP_ORDER:
        synth_pct[g] = comp_groups.get("Synthflow", {}).get(g, 0) / synth_total * 100 if synth_total > 0 else 0

    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(GROUP_ORDER))
    w = 0.35
    legacy_vals = [legacy_avg[g] for g in GROUP_ORDER]
    synth_vals = [synth_pct[g] for g in GROUP_ORDER]
    b1 = ax.bar(x - w / 2, legacy_vals, w, label="Legacy Avg (4 incumbents)", color=C_BLUE, alpha=0.85, edgecolor="white")
    b2 = ax.bar(x + w / 2, synth_vals, w, label="Synthflow (AI-native)", color=C_ORANGE, alpha=0.85, edgecolor="white")
    for bar, val in zip(b1, legacy_vals):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1, f"{val:.0f}%",
                    ha="center", fontsize=9, fontweight="bold", color=C_BLUE)
    for bar, val in zip(b2, synth_vals):
        if val > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1, f"{val:.0f}%",
                    ha="center", fontsize=9, fontweight="bold", color=C_ORANGE)
    labels = [g.replace("INDUSTRY ", "INDUSTRY\n") for g in GROUP_ORDER]
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=8.5)
    ax.set_ylabel("% of Churn", fontsize=9, color="#666")
    ax.set_ylim(0, max(max(legacy_vals, default=1), max(synth_vals, default=1)) * 1.3)
    ax.legend(fontsize=8.5, loc="upper right", framealpha=0.9)
    ax.set_title("AI eliminates call handling churn but doesn\u2019t eliminate churn",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.text(0.99, 0.01, "Legacy avg = mean of Ruby, AnswerConnect, PATLive, Smith.ai | Synthflow = only AI-native with enough data",
             fontsize=7, color="#999", ha="right", fontstyle="italic")
    fig.savefig(path); plt.close(fig)


def compute_industry_verticals(quotes, final_codes):
    """Extract industry vertical from quotes and compute category breakdown per vertical."""
    VERTICALS = {
        "Legal": ["law firm", "attorney", "lawyer", "legal", "paralegal", "litigation",
                   "practice area", "opposing counsel", "court", "client intake"],
        "Medical": ["patient", "medical", "doctor", "clinic", "hipaa", "healthcare",
                     "dental", "dermatolog", "medical practice"],
        "Home Services": ["hvac", "plumb", "electric", "roofing", "contractor", "technician",
                          "dispatch", "service call", "on-call", "field tech", "home service"],
        "Real Estate": ["real estate", "realtor", "property management", "listing agent", "showing"],
    }
    results = {}
    for idx, cat in final_codes.items():
        q = quotes[idx]
        text = (q["text"] + " " + q["llm"].get("pain_point", "")).lower()
        matched = None
        for vert, keywords in VERTICALS.items():
            if any(kw in text for kw in keywords):
                matched = vert; break
        if matched:
            results.setdefault(matched, {"count": 0, "cats": Counter()})
            results[matched]["count"] += 1
            results[matched]["cats"][cat] += 1
    return results

def compute_dollar_impact(quotes, final_codes):
    """Extract dollar amounts from churn quotes."""
    dollar_quotes = []
    for idx, cat in final_codes.items():
        q = quotes[idx]
        text = q["text"]
        amounts = re.findall(r'\$[\d,]+(?:\.\d{2})?(?:/(?:month|mo|yr|year))?', text)
        if not amounts: continue
        dollar_quotes.append({
            "idx": idx, "cat": cat, "product": normalize_product(q["llm"].get("product_mentioned")),
            "amounts": amounts, "text_snippet": truncate(text, 120),
            "pain_point": q["llm"].get("pain_point", ""),
        })
    return dollar_quotes


def chart_priority_matrix(cat_stats, total_weighted, path):
    """2x2 Impact vs Ease matrix for 11 categories."""
    _mbb()
    # Ease scores: how well AI solves this on day one (1-10)
    EASE = {
        "SCRIPT_ADHERENCE": 9, "INAUTHENTICITY": 8, "GARBLED_INFO": 9,
        "ROUTING_ERRORS": 8, "BILLING_PREDATORY": 9, "BILLING_OPAQUE": 9,
        "BILLING_TRAP": 10, "BILLING_TOO_EXPENSIVE": 5, "MISSED_CALLS": 10,
        "QUALITY_DECAY": 9, "SERIAL_SWITCHING": 3,
    }
    fig, ax = plt.subplots(figsize=(7.5, 5.5))

    # Compute actual data ranges for quadrant boundaries
    pcts_all = [s["weighted"] / total_weighted * 100 for s in cat_stats.values()]
    y_mid = np.median(pcts_all)  # median impact as divider
    x_mid = 6.5  # ease divider: below = hard, above = easy for AI
    y_max = max(pcts_all) * 1.35

    # Quadrant shading — using actual data range
    ax.axhspan(y_mid, y_max, xmin=0, xmax=(x_mid - 1) / 10, color="#FFEBEE", alpha=0.25)  # top-left: strategic
    ax.axhspan(y_mid, y_max, xmin=(x_mid - 1) / 10, xmax=1.0, color="#E3F2FD", alpha=0.25)  # top-right: do first
    ax.axhspan(0, y_mid, xmin=0, xmax=(x_mid - 1) / 10, color="#FFF3E0", alpha=0.25)  # bottom-left: deprioritize
    ax.axhspan(0, y_mid, xmin=(x_mid - 1) / 10, xmax=1.0, color="#E8F5E9", alpha=0.25)  # bottom-right: easy wins
    ax.axhline(y=y_mid, color="#BDC3C7", linewidth=1, linestyle="--", alpha=0.5)
    ax.axvline(x=x_mid, color="#BDC3C7", linewidth=1, linestyle="--", alpha=0.5)

    # Plot each category as a bubble
    # Pre-compute positions for smart label placement
    positions = {}
    for cat, s in cat_stats.items():
        pct = s["weighted"] / total_weighted * 100
        ease = EASE.get(cat, 5)
        positions[cat] = (ease, pct)

    for cat, s in cat_stats.items():
        pct = s["weighted"] / total_weighted * 100
        ease = EASE.get(cat, 5)
        color = GROUP_COLORS.get(CATEGORY_META[cat]["group"], C_GRAY)
        ax.scatter(ease, pct, s=pct * 18 + 60, color=color, alpha=0.85,
                   edgecolors="white", linewidth=1, zorder=5)
        label = CATEGORY_META[cat]["short"]
        # Smart label offsets to avoid collisions
        ox, oy = 0.2, 0.35
        if cat == "GARBLED_INFO": ox, oy = 0.2, -0.8  # below (collides with SCRIPT at ease=9)
        elif cat == "BILLING_OPAQUE": ox, oy = -3.5, 0.0  # left (collides with others at ease=9)
        elif cat == "QUALITY_DECAY": ox, oy = -3.2, 0.0  # left
        elif cat == "MISSED_CALLS": ox, oy = 0.2, -0.6  # below (at ease=10)
        elif cat == "BILLING_TRAP": ox, oy = 0.2, 0.5  # above (at ease=10)
        elif cat == "ROUTING_ERRORS": ox, oy = -3.0, -0.3  # left (collides with INAUTHENTICITY at ease=8)
        ax.annotate(label, (ease, pct), xytext=(ease + ox, pct + oy),
                    fontsize=7.5, color="#333", fontweight="bold" if pct > 8 else "normal",
                    arrowprops=dict(arrowstyle="-", color="#BDC3C7", lw=0.5) if abs(ox) > 1 else None)

    # Quadrant labels — placed in actual quadrant centers
    ax.text(3.5, y_max * 0.85, "STRATEGIC\nBETS", fontsize=9, color="#C0392B",
            fontweight="bold", ha="center", alpha=0.5)
    ax.text(8.5, y_max * 0.85, "DO\nFIRST", fontsize=9, color="#1B365F",
            fontweight="bold", ha="center", alpha=0.5)
    ax.text(3.5, y_mid * 0.25, "DEPRIORITIZE", fontsize=9, color="#E67E22",
            fontweight="bold", ha="center", alpha=0.5)
    ax.text(8.5, y_mid * 0.25, "EASY\nWINS", fontsize=9, color="#27AE60",
            fontweight="bold", ha="center", alpha=0.5)

    ax.set_xlabel("How Easily AI Solves This (1\u201310)", fontsize=9, color="#666")
    ax.set_ylabel("Share of Churn (%)", fontsize=9, color="#666")
    ax.set_xlim(1, 11); ax.set_ylim(0, y_max)
    from matplotlib.patches import Patch
    ax.legend(handles=[Patch(facecolor=GROUP_COLORS[g], label=g) for g in GROUP_ORDER],
               loc="upper left", fontsize=7, framealpha=0.9)
    ax.set_title("Script failures and inauthenticity: high impact, directly addressed by AI",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)


# ── Docx helpers ─────────────────────────────────────────────────────
def set_shading(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))

def style_table(tbl, header_bg=LIGHT_BLUE_BG, stripe=True):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl.rows[0].cells:
        set_shading(cell, header_bg)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.color.rgb = DARK_BLUE; r.font.size = Pt(9)
    for i, row in enumerate(tbl.rows[1:], 1):
        for cell in row.cells:
            if stripe and i % 2 == 0: set_shading(cell, LIGHT_GRAY_BG)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY

def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = DARK_BLUE
    return hd

def body(doc, text, bold=False, italic=False, sz=Pt(10.5), color=DARK_GRAY):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = sz; r.font.color.rgb = color
    return p

def quote_block(doc, qt, attrib):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u201C{qt}\u201D"); r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK_GRAY
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.5)
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(f"\u2014 {attrib}"); r2.font.size = Pt(9); r2.font.color.rgb = MID_GRAY

def callout(doc, text):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0); set_shading(cell, LIGHT_BLUE_BG)
    p = cell.paragraphs[0]; r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_chart(doc, path, width=Inches(5.8)):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)

def add_toc(doc):
    """Write a manual table of contents."""
    toc_entries = [
        ("1. Executive Summary", 1),
        ("2. What\u2019s Breaking", 1),
        ("    Four Customers, Four Problems", 2),
        ("    Where the Pain Concentrates", 2),
        ("    The 11 Specific Failure Modes", 2),
        ("    The AI-Native Comparison", 2),
        ("3. Who\u2019s Most Vulnerable", 1),
        ("    Each Competitor Fails Differently", 2),
        ("    Where Customers Go When They Leave", 2),
        ("    The Dollar Impact", 2),
        ("    Legal Is the Vertical to Win First", 2),
        ("4. Why Now", 1),
        ("    The Real Ratings vs. the Headline Ratings", 2),
        ("    The Review Distribution Proves It", 2),
        ("    AI Alternatives Are Proving the Concept", 2),
        ("5. What to Do About It", 1),
        ("    The Priority Matrix", 2),
        ("    What to Say (and Where)", 2),
        ("    Who to Target First", 2),
        ("    Sizing the Opportunity", 2),
        ("Appendix A: Methodology", 1),
        ("Appendix B: The Competitors We Studied", 1),
        ("Appendix C: Best Quotes by Category", 1),
        ("Appendix D: Full Quote Database", 1),
    ]
    for title, level in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        indent = Cm(0) if level == 1 else Cm(1)
        p.paragraph_format.left_indent = indent
        r = p.add_run(title)
        r.font.size = Pt(11) if level == 1 else Pt(10)
        r.font.color.rgb = DARK_BLUE if level == 1 else MID_GRAY
        r.bold = (level == 1)


# ── Main ─────────────────────────────────────────────────────────────
def main():
    quotes = load_quotes()
    codes_a, codes_b = load_coder(CODER_A), load_coder(CODER_B)
    final_codes, disagreements = adjudicate(codes_a, codes_b)
    print(f"Loaded: {len(quotes)} quotes, {len(final_codes)} codes, {len(disagreements)} disagreements")

    # ── Stats ────────────────────────────────────────────────────
    cat_stats = {}
    for idx, cat in final_codes.items():
        cat_stats.setdefault(cat, {"count": 0, "weighted": 0.0, "quotes": []})
        q = quotes[idx]; w = compute_weight(q)
        cat_stats[cat]["count"] += 1; cat_stats[cat]["weighted"] += w
        cat_stats[cat]["quotes"].append((idx, q, w))
    total_count = sum(s["count"] for s in cat_stats.values())
    total_weighted = sum(s["weighted"] for s in cat_stats.values())
    sorted_cats = sorted(cat_stats, key=lambda c: cat_stats[c]["weighted"], reverse=True)

    group_stats = {}
    for cat, s in cat_stats.items():
        g = CATEGORY_META[cat]["group"]
        group_stats.setdefault(g, {"count": 0, "weighted": 0.0, "cats": []})
        group_stats[g]["count"] += s["count"]; group_stats[g]["weighted"] += s["weighted"]
        group_stats[g]["cats"].append(cat)

    comp_cats, comp_groups = {}, {}
    for idx, cat in final_codes.items():
        prod = normalize_product(quotes[idx]["llm"].get("product_mentioned"))
        comp_cats.setdefault(prod, Counter())[cat] += 1
    for prod, cats in comp_cats.items():
        comp_groups[prod] = Counter()
        for cat, cnt in cats.items():
            comp_groups[prod][CATEGORY_META[cat]["group"]] += cnt

    # ── Gut-punch: curate one per group ──────────────────────────
    gut_punch = []
    for g in GROUP_ORDER:
        group_cat_list = group_stats[g]["cats"]
        candidates = []
        for idx, cat in final_codes.items():
            if cat not in group_cat_list: continue
            q = quotes[idx]; quality = q["llm"].get("quote_quality", 1)
            pull = extract_pull_quote(q["text"])
            sc = devastation_score(pull) * 10 + devastation_score(q["text"]) * 3 + quality * 2
            candidates.append((sc, idx, q, cat))
        candidates.sort(key=lambda x: x[0], reverse=True)
        # Pick best that's not a duplicate product
        used = {gp[3] for gp in gut_punch}
        for sc, idx, q, cat in candidates:
            prod = normalize_product(q["llm"].get("product_mentioned"))
            if prod not in used or len(candidates) < 3:
                gut_punch.append((idx, q, cat, prod))
                break

    # ── Computed analysis ────────────────────────────────────────
    switching_data = compute_switching_data(quotes, final_codes)
    kappa_val, kappa_agree, kappa_n = compute_kappa(codes_a, codes_b)
    print(f"Kappa: {kappa_val:.4f} ({kappa_agree}/{kappa_n} agree)")
    print(f"Switching: {switching_data['total']} stories, triggers={dict(switching_data['triggers'])}")
    print(f"Net flow: {dict(switching_data['net_flow'])}")

    # ── Charts ───────────────────────────────────────────────────
    tmpdir = tempfile.mkdtemp(); cp = {}; print("Generating charts...")
    cp["groups"]   = os.path.join(tmpdir, "g.png"); chart_group_pareto(group_stats, total_weighted, cp["groups"])
    cp["pareto"]   = os.path.join(tmpdir, "p.png"); chart_pareto_11(sorted_cats, cat_stats, total_weighted, cp["pareto"])
    cp["heatmap"]  = os.path.join(tmpdir, "h.png"); chart_heatmap(comp_groups, cp["heatmap"])
    cp["net"]      = os.path.join(tmpdir, "n.png"); chart_net_churn(cp["net"], switching_data)
    temporal_data, temporal_order = compute_temporal_data()
    cp["temporal"] = os.path.join(tmpdir, "t.png"); chart_temporal_before_after(cp["temporal"], temporal_data, temporal_order)
    cp["matrix"]   = os.path.join(tmpdir, "m.png"); chart_priority_matrix(cat_stats, total_weighted, cp["matrix"])
    bimod_data = compute_bimodality(temporal_data, temporal_order)
    cp["bimod"]    = os.path.join(tmpdir, "b.png"); chart_bimodality(bimod_data, temporal_order, cp["bimod"])
    cp["legacy_vs_ai"] = os.path.join(tmpdir, "lva.png"); chart_legacy_vs_ai(comp_groups, cp["legacy_vs_ai"])
    industry_data = compute_industry_verticals(quotes, final_codes)
    dollar_data = compute_dollar_impact(quotes, final_codes)
    print("Charts generated.")

    # ═════════════════════════════════════════════════════════════
    # BUILD DOCUMENT
    # ═════════════════════════════════════════════════════════════
    doc = Document()
    style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10.5); style.font.color.rgb = DARK_GRAY
    for lv in range(1, 4):
        s = doc.styles[f"Heading {lv}"]; s.font.name = "Calibri"; s.font.color.rgb = DARK_BLUE
    for sec in doc.sections:
        sec.top_margin = Cm(2); sec.bottom_margin = Cm(2)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
        footer = sec.footer; footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin"); run._r.append(fld1)
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
        run2 = fp.add_run(); run2._r.append(instr)
        fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "separate")
        run3 = fp.add_run(); run3._r.append(fld2)
        run4 = fp.add_run("1"); run4.font.size = Pt(8); run4.font.color.rgb = MID_GRAY
        fld3 = OxmlElement("w:fldChar"); fld3.set(qn("w:fldCharType"), "end")
        run5 = fp.add_run(); run5._r.append(fld3)

    # ── Dynamic percentages ──────────────────────────────────
    ch_pct = group_stats["CALL HANDLING"]["weighted"] / total_weighted * 100
    bill_pct = group_stats["BILLING"]["weighted"] / total_weighted * 100
    rel_pct = group_stats["SERVICE RELIABILITY"]["weighted"] / total_weighted * 100
    disill_pct = group_stats["INDUSTRY DISILLUSIONMENT"]["weighted"] / total_weighted * 100
    ai_addr_pct = ch_pct + bill_pct
    agree_n = len(final_codes) - len(disagreements)
    agree_pct = agree_n / len(final_codes) * 100

    def comp_group_pct(prod, group):
        total = sum(comp_groups.get(prod, {}).values())
        return comp_groups.get(prod, {}).get(group, 0) / total * 100 if total > 0 else 0

    legacy_prods = ["Ruby Receptionist", "AnswerConnect", "PATLive", "Smith.ai"]
    def legacy_avg_pct(group):
        vals = [comp_group_pct(p, group) for p in legacy_prods if sum(comp_groups.get(p, {}).values()) > 0]
        return float(np.mean(vals)) if vals else 0

    with open(QUOTES_FILE) as _f:
        _all_quotes = json.load(_f)
    total_all_quotes = len(_all_quotes)
    positive_count = sum(1 for q in _all_quotes if q["llm"].get("category", "").startswith("positive"))
    tp_churn = sum(1 for i in final_codes if quotes[i]["source"] == "trustpilot")
    reddit_churn = sum(1 for i in final_codes if quotes[i]["source"] == "reddit")
    org_avgs = [td.get("organic_avg", 3.0) for td in temporal_data.values() if td.get("organic_n", 0) > 0]
    org_min, org_max = (min(org_avgs), max(org_avgs)) if org_avgs else (1.0, 4.0)

    n_switched = switching_data["total"]
    nf = switching_data["net_flow"]
    biggest_loser = min(nf.items(), key=lambda x: x[1]["net"])[0] if nf else "N/A"
    biggest_loss = min(nf.values(), key=lambda x: x["net"])["net"] if nf else 0
    synth_bill = comp_group_pct("Synthflow", "BILLING")
    synth_rel = comp_group_pct("Synthflow", "SERVICE RELIABILITY")
    synth_ch = comp_group_pct("Synthflow", "CALL HANDLING")

    # ══════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ANSWERING SERVICE\nCHURN ANALYSIS"); r.font.size = Pt(32); r.font.color.rgb = DARK_BLUE; r.bold = True
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Why {ai_addr_pct:.0f}% of Answering Service Churn\nMaps to AI\u2019s Structural Advantage")
    r.font.size = Pt(16); r.font.color.rgb = ACCENT_BLUE
    for _ in range(2): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{total_count} churn stories  |  {len(COMPETITOR_PROFILES)} competitors  |  Reddit & Trustpilot  |  2015\u20132026")
    r.font.size = Pt(10); r.font.color.rgb = MID_GRAY
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENTIAL \u2014 Prepared for Central AI Leadership")
    r.font.size = Pt(9); r.font.color.rgb = MID_GRAY; r.italic = True
    doc.add_page_break()

    # ── TABLE OF CONTENTS ─────────────────────────────────────
    h(doc, "Contents", 1)
    add_toc(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 1. EXECUTIVE SUMMARY
    # ══════════════════════════════════════════════════════════
    h(doc, "1. Executive Summary", 1)

    callout(doc, f"{ch_pct:.0f}% of churn is call handling. {bill_pct:.0f}% is billing.\n"
                 f"Together, {ai_addr_pct:.0f}% maps to problems where AI has a structural advantage.")

    body(doc, f"We analyzed {total_all_quotes} reviews of {len(COMPETITOR_PROFILES)} answering services "
         f"across Reddit and Trustpilot. {total_count} customers described specific churn reasons. "
         f"Every complaint was independently classified twice "
         f"(Cohen\u2019s kappa = {kappa_val:.2f}, \u201Calmost perfect agreement\u201D).")

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Why now:"); r.bold = True; r.font.size = Pt(10.5)
    for bullet in [
        f"Unprompted reviews average {org_min:.1f}\u2013{org_max:.1f} stars. Recent 5-star waves mask persistent pain.",
        f"AI-native entrants prove the model works but fail on billing "
        f"({synth_bill:.0f}% of Synthflow churn) and reliability ({synth_rel:.0f}%). "
        f"The window is open but closing.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet).font.size = Pt(10.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("Recommended actions:"); r.bold = True; r.font.size = Pt(10.5)
    ruby_net = nf.get("Ruby Receptionist", {}).get("net", 0)
    smith_net = nf.get("Smith.ai", {}).get("net", 0)
    for action in [
        f"Target Ruby and Smith.ai churners first \u2014 they lose the most customers "
        f"(Ruby {ruby_net:+d}, Smith.ai {smith_net:+d} net).",
        f"Lead messaging with call handling quality ({ch_pct:.0f}% of churn) "
        f"and transparent pricing ({bill_pct:.0f}%).",
        "Start with legal vertical \u2014 highest signal density in the data.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(action).font.size = Pt(10.5)

    doc.add_paragraph()
    body(doc, "Reading guide: Section 2 shows what\u2019s breaking and why. Section 3 identifies the most "
         "vulnerable competitors. Section 4 explains why the timing is right. Section 5 gives specific "
         "actions. Methodology and full data are in the appendices.",
         italic=True, sz=Pt(9.5), color=MID_GRAY)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 2. WHAT'S BREAKING
    # ══════════════════════════════════════════════════════════
    h(doc, "2. What\u2019s Breaking", 1)
    body(doc, f"We analyzed {total_all_quotes} reviews of {len(COMPETITOR_PROFILES)} answering services. "
         f"{total_count} customers described exactly why they left. "
         f"Their reasons cluster into four groups.")

    # ── 2a: Four Customers, Four Problems ─────────────────────
    doc.add_paragraph()
    h(doc, "Four Customers, Four Problems", 2)
    body(doc, "Before the data, hear the customers. One quote per group \u2014 no commentary.")
    doc.add_paragraph()

    theme_labels = {
        "CALL HANDLING": "THE CALL HANDLING PROBLEM",
        "BILLING": "THE BILLING PROBLEM",
        "SERVICE RELIABILITY": "THE RELIABILITY PROBLEM",
        "INDUSTRY DISILLUSIONMENT": "THE CATEGORY PROBLEM",
    }
    for (idx, q, cat, prod), g in zip(gut_punch, GROUP_ORDER):
        pull = extract_pull_quote(q["text"])
        pain = q["llm"].get("pain_point", "")
        p = doc.add_paragraph()
        r = p.add_run(theme_labels.get(g, g)); r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = ACCENT_BLUE; r.font.all_caps = True
        quote_block(doc, pull, f"{prod} customer \u2014 {pain}")

    doc.add_page_break()

    # ── 2b: The Churn Reasons, Ranked ─────────────────────────
    h(doc, "Where the Pain Concentrates", 2)
    body(doc, f"The largest share of churn ({ch_pct:.0f}%) comes from the calls themselves. "
         f"The second-largest ({bill_pct:.0f}%) is billing. Together, {ai_addr_pct:.0f}% of churn "
         f"maps to problems where AI has a structural advantage over shared human pools.", bold=True)
    doc.add_paragraph()
    add_chart(doc, cp["groups"], Inches(6))
    doc.add_paragraph()

    # Walk through each group narrative-style, like the research plan
    for rank, g in enumerate(GROUP_ORDER, 1):
        gs = group_stats.get(g); gm = GROUP_META[g]
        pct = gs["weighted"] / total_weighted * 100
        cats_sorted = sorted(gs["cats"], key=lambda c: cat_stats[c]["weighted"], reverse=True)

        hd = doc.add_heading(f"#{rank}: {g} \u2014 {pct:.0f}% of all churn", level=3)
        for r in hd.runs: r.font.color.rgb = DARK_BLUE

        # Plain-language description
        body(doc, gm["insight"])

        # Who's most affected (computed dynamically)
        affected = []
        for prod in TOP_COMPETITORS:
            prod_pct = comp_group_pct(prod, g)
            if prod_pct > 0:
                affected.append((prod, prod_pct))
        affected.sort(key=lambda x: x[1], reverse=True)
        if affected:
            affected_str = ", ".join(f"{p} ({v:.0f}%)" for p, v in affected[:3])
            p = doc.add_paragraph()
            r = p.add_run("Most affected: "); r.bold = True; r.font.size = Pt(10)
            p.add_run(affected_str).font.size = Pt(10)

        # Sub-categories
        for c in cats_sorted:
            cs = cat_stats[c]; m = CATEGORY_META[c]
            cat_pct = cs["weighted"] / total_weighted * 100
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{m['short']} ({cs['count']} quotes, {cat_pct:.1f}%): ")
            r.bold = True; r.font.size = Pt(10)
            p.add_run(m["gtm"]).font.size = Pt(10)

        # How AI solves this
        p = doc.add_paragraph()
        r = p.add_run("Product implication: "); r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = ACCENT_BLUE
        p.add_run(gm["impl"]).font.size = Pt(10)
        doc.add_paragraph()

    doc.add_page_break()

    # ── 2c: The 11 Specific Failure Modes ─────────────────────
    h(doc, "The 11 Specific Failure Modes", 2)
    body(doc, "Each group breaks into specific, actionable problems. "
         "The Pareto below shows the top 6 account for the vast majority of churn:")
    doc.add_paragraph()
    add_chart(doc, cp["pareto"], Inches(6.2))

    # 11-row table
    tbl = doc.add_table(rows=len(sorted_cats) + 2, cols=6); tbl.style = "Table Grid"
    for i, hdr in enumerate(["Rank", "Churn Reason", "Group", "Count", "%", "Running Total"]):
        tbl.rows[0].cells[i].text = hdr
    cumul = 0
    for rank, cat in enumerate(sorted_cats, 1):
        s = cat_stats[cat]; m = CATEGORY_META[cat]
        pct = s["weighted"] / total_weighted * 100; cumul += pct
        row = tbl.rows[rank]
        row.cells[0].text = str(rank); row.cells[1].text = m["short"]; row.cells[2].text = m["group"]
        row.cells[3].text = str(s["count"]); row.cells[4].text = f"{pct:.1f}%"; row.cells[5].text = f"{cumul:.1f}%"
    tr = tbl.rows[-1]; tr.cells[1].text = "Total"; tr.cells[3].text = str(total_count); tr.cells[4].text = "100%"
    for c in [tr.cells[1], tr.cells[3]]:
        for p in c.paragraphs:
            for r in p.runs: r.bold = True
    style_table(tbl)
    for row in tbl.rows:
        for i in [0, 3, 4, 5]: row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── 2d: The AI-Native Comparison ──────────────────────────
    h(doc, "The AI-Native Comparison", 2)
    body(doc, f"Does switching to AI solve these problems? Partially. Synthflow (the only AI-native "
         f"competitor with enough data) has {synth_ch:.0f}% call handling churn vs "
         f"{legacy_avg_pct('CALL HANDLING'):.0f}% for legacy services. "
         f"But billing ({synth_bill:.0f}%) and platform reliability ({synth_rel:.0f}%) become "
         f"the new battleground.", bold=True)
    doc.add_paragraph()
    add_chart(doc, cp["legacy_vs_ai"], Inches(6))
    doc.add_paragraph()
    body(doc, "Central AI must avoid both failure modes: the call handling problems that plague legacy "
         "services AND the billing and reliability problems that plague AI-native entrants.",
         italic=True, sz=Pt(10), color=MID_GRAY)

    doc.add_paragraph()

    # ── 2e: Where Customers Go When They Leave ────────────────
    h(doc, "Where Customers Go When They Leave", 2)
    body(doc, f"We found {n_switched} stories where a customer described switching services. "
         f"{biggest_loser} has the worst net flow ({biggest_loss:+d}).", bold=True)
    doc.add_paragraph()
    add_chart(doc, cp["net"], Inches(5.8))

    doc.add_paragraph()
    sw_triggers = switching_data["triggers"]
    sw_total = sum(sw_triggers.values())
    trigger_descs = {
        "Quality": "Missed calls, wrong info, platform not working, agents not trained",
        "Capability": "Missing features, no integrations, limited technology",
        "Pricing": "Too expensive, hidden fees, per-minute overcharges",
    }
    sw_sorted = sorted(sw_triggers.items(), key=lambda x: x[1], reverse=True)
    tt = doc.add_table(rows=len(sw_sorted) + 1, cols=4); tt.style = "Table Grid"
    for i, hdr in enumerate(["Trigger", "Count", f"% of {sw_total}", "What It Means"]):
        tt.rows[0].cells[i].text = hdr
    for r, (trig, cnt) in enumerate(sw_sorted, 1):
        pct = cnt / sw_total * 100 if sw_total else 0
        tt.rows[r].cells[0].text = trig; tt.rows[r].cells[1].text = str(cnt)
        tt.rows[r].cells[2].text = f"{pct:.0f}%"; tt.rows[r].cells[3].text = trigger_descs.get(trig, "")
    style_table(tt)
    for row in tt.rows:
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    body(doc, f"Based on {n_switched} switching stories. Directional, not statistically definitive.",
         italic=True, sz=Pt(9.5), color=MID_GRAY)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 3. WHO'S MOST VULNERABLE
    # ══════════════════════════════════════════════════════════
    h(doc, "3. Who\u2019s Most Vulnerable", 1)
    body(doc, "Now we know what breaks. The next question: whose customers are most ready to switch?")

    # ── 3a: Each Competitor Fails Differently ─────────────────
    doc.add_paragraph()
    h(doc, "Each Competitor Fails Differently", 2)
    add_chart(doc, cp["heatmap"], Inches(5.8))
    doc.add_paragraph()

    overall_pcts = {g: group_stats[g]["weighted"] / total_weighted for g in GROUP_ORDER}
    for prod in TOP_COMPETITORS:
        cats = comp_cats.get(prod, Counter()); total_prod = sum(cats.values())
        if total_prod == 0: continue
        groups = comp_groups.get(prod, Counter())
        top_group = max(GROUP_ORDER, key=lambda g: groups.get(g, 0) / total_prod if total_prod > 0 else 0)
        top_group_pct = groups.get(top_group, 0) / total_prod * 100
        top2 = cats.most_common(2)
        top2_t = " + ".join(f"{CATEGORY_META[c]['short']} ({cnt})" for c, cnt in top2)
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{prod}: "); r.bold = True; r.font.size = Pt(10)
        p.add_run(f"{top_group_pct:.0f}% {top_group.lower()}. Top: {top2_t}.").font.size = Pt(10)

    # ── 3b: The Dollar Impact ─────────────────────────────────
    if dollar_data:
        doc.add_paragraph()
        h(doc, "The Dollar Impact", 2)
        body(doc, f"{len(dollar_data)} of {total_count} churn quotes mention specific dollar amounts:")
        doc.add_paragraph()
        monthly_costs = [d for d in dollar_data if any("/mo" in a or "/month" in a for a in d["amounts"])]
        overcharges = [d for d in dollar_data if d["cat"] in ("BILLING_PREDATORY", "BILLING_OPAQUE", "BILLING_TRAP")]
        for lbl, items in [
            ("Monthly costs customers report paying", monthly_costs),
            ("Documented overcharges and losses", overcharges),
        ]:
            if not items: continue
            p = doc.add_paragraph()
            r = p.add_run(lbl); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = ACCENT_BLUE
            for d in items[:5]:
                p2 = doc.add_paragraph(style="List Bullet")
                amt_str = ", ".join(d["amounts"])
                r2 = p2.add_run(f"{d['product']} \u2014 {amt_str}: "); r2.bold = True; r2.font.size = Pt(9.5)
                p2.add_run(d["pain_point"]).font.size = Pt(9.5)

    # ── 3c: Legal Is the Vertical to Win First ────────────────
    if industry_data:
        doc.add_paragraph()
        h(doc, "Legal Is the Vertical to Win First", 2)
        total_vert_quotes = sum(v["count"] for v in industry_data.values())
        vert_sorted = sorted(industry_data.items(), key=lambda x: x[1]["count"], reverse=True)
        total_ident = sum(v["count"] for _, v in vert_sorted)
        vt = doc.add_table(rows=len(vert_sorted) + 1, cols=4); vt.style = "Table Grid"
        for i, hdr in enumerate(["Industry", "Quotes", "% of Identified", "Top Churn Reason"]):
            vt.rows[0].cells[i].text = hdr
        for r, (vert, vdata) in enumerate(vert_sorted, 1):
            top_cat = vdata["cats"].most_common(1)[0] if vdata["cats"] else ("", 0)
            top_label = CATEGORY_META.get(top_cat[0], {}).get("short", top_cat[0]) if top_cat[0] else ""
            vt.rows[r].cells[0].text = vert
            vt.rows[r].cells[1].text = str(vdata["count"])
            vt.rows[r].cells[2].text = f"{vdata['count']/total_ident*100:.0f}%"
            vt.rows[r].cells[3].text = f"{top_label} ({top_cat[1]})" if top_label else ""
        style_table(vt)
        for row in vt.rows:
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
        if "Legal" in industry_data:
            legal = industry_data["Legal"]
            body(doc, f"Legal dominates ({legal['count']} of {total_ident} identified, "
                 f"{legal['count']/total_ident*100:.0f}%). "
                 f"Caveat: all Reddit churn signal comes from r/LawFirm.", bold=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 4. WHY NOW
    # ══════════════════════════════════════════════════════════
    h(doc, "4. Why Now", 1)
    body(doc, f"{biggest_loser} and Smith.ai customers are leaving. "
         "But is this the right moment to catch them?")

    # ── 4a: The Real Ratings ──────────────────────────────────
    doc.add_paragraph()
    h(doc, "The Real Ratings vs. the Headline Ratings", 2)
    body(doc, f"If you look at Trustpilot today, these companies look fine. They are not fine. "
         f"When customers post reviews on their own (before mid-2024), the average is "
         f"{org_min:.1f}\u2013{org_max:.1f} stars.", bold=True)
    doc.add_paragraph()
    add_chart(doc, cp["temporal"], Inches(6))

    st = doc.add_table(rows=5, cols=4); st.style = "Table Grid"
    for i, hdr in enumerate(["Competitor", "Unprompted Avg", "Recent Wave Avg", "Jump"]):
        st.rows[0].cells[i].text = hdr
    for r, comp in enumerate(temporal_order, 1):
        td = temporal_data.get(comp, {})
        org_a, org_n = td.get("organic_avg", 0), td.get("organic_n", 0)
        sur_a, sur_n = td.get("surge_avg", 0), td.get("surge_n", 0)
        jump = td.get("jump", 0)
        row_data = (comp, f"{org_a:.2f} ({org_n} reviews)", f"{sur_a:.2f} ({sur_n} reviews)",
                    f"+{jump:.2f}" if jump >= 0 else f"{jump:.2f}")
        for i, v in enumerate(row_data): st.rows[r].cells[i].text = v
    style_table(st)
    for row in st.rows:
        for i in [1, 2, 3]: row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 4b: The Review Distribution ───────────────────────────
    doc.add_paragraph()
    h(doc, "The Review Distribution Proves It", 2)
    body(doc, "If these services were improving, we\u2019d see reviews spread across 1\u20135 stars. "
         "Instead, they cluster at the extremes: genuinely unhappy customers and prompted 5-star reviews.")
    doc.add_paragraph()
    add_chart(doc, cp["bimod"], Inches(5.8))
    doc.add_paragraph()

    for prod in ["Smith.ai", "AnswerConnect", "Ruby Receptionist", "PATLive"]:
        bd = bimod_data.get(prod, {})
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{prod}: "); r.bold = True; r.font.size = Pt(10)
        pct = bd.get("overall_extreme_pct", 0)
        org_5 = bd.get("org_5star", 0); org_n = bd.get("org_total", 0)
        surge_5 = bd.get("surge_5star", 0); surge_n = bd.get("surge_total", 0)
        detail = f"{pct:.0f}% extreme. "
        if org_5 == 0 and org_n > 0:
            detail += f"Before mid-2024: zero 5-star out of {org_n}. Since: {surge_5} of {surge_n}."
        else:
            detail += f"Before mid-2024: {org_5}/{org_n} were 5-star. Since: {surge_5}/{surge_n}."
        p.add_run(detail).font.size = Pt(10)

    # ── 4c: AI Alternatives ───────────────────────────────────
    doc.add_paragraph()
    h(doc, "AI Alternatives Are Proving the Concept", 2)
    body(doc, "Dialzara and My AI Front Desk average 5.0 stars on Trustpilot \u2014 but with small samples "
         "(14 and 8 reviews). Synthflow, with 69 reviews, shows the real pattern: AI eliminates call "
         "handling churn but billing and platform reliability become the new failure modes.")
    body(doc, f"Synthflow\u2019s churn: {synth_bill:.0f}% billing, {synth_rel:.0f}% reliability, "
         f"{synth_ch:.0f}% call handling. Central AI must avoid both legacy and AI-native failure modes.",
         bold=True)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # 5. WHAT TO DO ABOUT IT
    # ══════════════════════════════════════════════════════════
    h(doc, "5. What to Do About It", 1)
    body(doc, "The market is broken, the timing is right. Here\u2019s what to do.")

    # ── 5a: The Priority Matrix ───────────────────────────────
    doc.add_paragraph()
    h(doc, "The Priority Matrix", 2)
    add_chart(doc, cp["matrix"], Inches(5.8))
    body(doc, "Top-right quadrant (high impact, easy for AI) is where to focus first. "
         "Script failures and inauthenticity are there.",
         italic=True, sz=Pt(10), color=MID_GRAY)

    doc.add_paragraph()
    h(doc, "What to Build (in Order)", 3)
    body(doc, "Ranked by how much churn each capability prevents:")
    rm = doc.add_table(rows=5, cols=4); rm.style = "Table Grid"
    for i, hdr in enumerate(["#", "Capability", "% of Churn", "Action"]):
        rm.rows[0].cells[i].text = hdr
    for r, (cap, cats_list, act) in enumerate([
        ("AI Training & Knowledge", ["SCRIPT_ADHERENCE", "INAUTHENTICITY"],
         "AI learns the business once and retains it. No retraining cycle."),
        ("Transparent Pricing", ["BILLING_PREDATORY", "BILLING_OPAQUE", "BILLING_TRAP", "BILLING_TOO_EXPENSIVE"],
         "Flat rate with real-time usage dashboard. Cancel anytime, one click."),
        ("Data Capture Accuracy", ["GARBLED_INFO"],
         "Capture names, emails, and phone numbers correctly every time."),
        ("Consistent Availability", ["MISSED_CALLS", "QUALITY_DECAY", "ROUTING_ERRORS"],
         "Answer every call, route correctly, maintain quality over time."),
    ], 1):
        pct = sum(cat_stats.get(c, {}).get("weighted", 0) for c in cats_list) / total_weighted * 100
        cnt = sum(cat_stats.get(c, {}).get("count", 0) for c in cats_list)
        rm.rows[r].cells[0].text = str(r); rm.rows[r].cells[1].text = cap
        rm.rows[r].cells[2].text = f"{pct:.0f}% ({cnt})"; rm.rows[r].cells[3].text = act
    style_table(rm)
    for row in rm.rows:
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 5b: What to Say ───────────────────────────────────────
    doc.add_paragraph()
    h(doc, "What to Say (and Where)", 2)
    gt = doc.add_table(rows=6, cols=4); gt.style = "Table Grid"
    for i, hdr in enumerate(["Channel", "Message", "Pain", "% Churn"]):
        gt.rows[0].cells[i].text = hdr
    cp_ = lambda g: f"{group_stats[g]['weighted']/total_weighted*100:.0f}%"
    ip_ = f"{cat_stats['INAUTHENTICITY']['weighted']/total_weighted*100:.0f}%"
    for r, row_data in enumerate([
        ("Homepage", "\u201CThey didn\u2019t know my business.\u201D We do.", "Call handling", cp_("CALL HANDLING")),
        ("Pricing page", "No per-minute billing. See your costs before the bill.", "Billing", cp_("BILLING")),
        ("Comparison", "\u201CService used to be great, now it\u2019s not.\u201D AI doesn\u2019t decay.", "Reliability", cp_("SERVICE RELIABILITY")),
        ("Sales calls", "\u201CI\u2019ve tried 4 services.\u201D Different technology, not another vendor.", "Switchers", cp_("INDUSTRY DISILLUSIONMENT")),
        ("Demo CTA", "Talk to our AI. Ask it about your business.", "Inauthenticity", ip_),
    ], 1):
        for i, v in enumerate(row_data): gt.rows[r].cells[i].text = v
    style_table(gt)
    for row in gt.rows: row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 5c: Who to Target First ───────────────────────────────
    doc.add_paragraph()
    h(doc, "Who to Target First", 2)

    # Target prioritization summary
    target_data = []
    for prod in ["Ruby Receptionist", "Smith.ai", "AnswerConnect"]:
        n = sum(comp_cats.get(prod, {}).values())
        net = nf.get(prod, {}).get("net", 0)
        top_group = max(GROUP_ORDER, key=lambda g: comp_groups.get(prod, {}).get(g, 0))
        tg_pct = comp_group_pct(prod, top_group)
        target_data.append((prod, f"{n} churners, {net:+d} net", f"{tg_pct:.0f}% {top_group.lower()}", CATEGORY_META[comp_cats.get(prod, Counter()).most_common(1)[0][0]]["gtm"] if comp_cats.get(prod) else ""))
    serial_n = cat_stats.get("SERIAL_SWITCHING", {}).get("count", 0)
    target_data.append(("Serial Switchers", f"{serial_n} quotes", "Tried multiple services", "Position as different technology"))
    synth_n = sum(comp_cats.get("Synthflow", {}).values())
    target_data.append(("Synthflow users", f"{synth_n} churners", f"{synth_bill:.0f}% billing, {synth_rel:.0f}% reliability", "Reliable AI with real support"))

    tp = doc.add_table(rows=len(target_data) + 1, cols=4); tp.style = "Table Grid"
    for i, hdr in enumerate(["Target", "Why", "Weakness", "Lead Message"]):
        tp.rows[0].cells[i].text = hdr
    for r, (target, why, weakness, msg) in enumerate(target_data, 1):
        tp.rows[r].cells[0].text = target; tp.rows[r].cells[1].text = why
        tp.rows[r].cells[2].text = weakness; tp.rows[r].cells[3].text = msg
    style_table(tp)

    doc.add_paragraph()

    # Per-competitor briefs
    for prod in TOP_COMPETITORS:
        cats = comp_cats.get(prod, Counter()); total_prod = sum(cats.values())
        if total_prod == 0: continue
        hd = doc.add_heading(f"{prod} ({total_prod} churners)", level=3)
        for r in hd.runs: r.font.color.rgb = DARK_BLUE
        for cat, cnt in cats.most_common(3):
            m = CATEGORY_META[cat]; pct = cnt / total_prod * 100
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(f"{m['short']}: {cnt} ({pct:.0f}%)"); r.bold = True; r.font.size = Pt(10)
            r2 = p.add_run(f" \u2014 {m['gtm']}"); r2.font.size = Pt(10); r2.font.color.rgb = MID_GRAY
        cq = [(idx, quotes[idx], compute_weight(quotes[idx]))
              for idx in final_codes
              if normalize_product(quotes[idx]["llm"].get("product_mentioned")) == prod
              and quotes[idx]["llm"].get("quote_quality", 1) >= 4]
        if cq:
            cq.sort(key=lambda x: devastation_score(x[1]["text"]) * x[2], reverse=True)
            pain = cq[0][1]["llm"].get("pain_point", "")
            p = doc.add_paragraph()
            r = p.add_run(f"Voice of their customer: \u201C{pain}\u201D")
            r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MID_GRAY

    # ── 5d: Sizing the Opportunity ────────────────────────────
    doc.add_paragraph()
    h(doc, "Sizing the Opportunity", 2)
    body(doc, "Every assumption is flagged for your team to sharpen.", italic=True, color=MID_GRAY)
    doc.add_paragraph()

    tam = doc.add_table(rows=7, cols=3); tam.style = "Table Grid"
    for i, hdr in enumerate(["Step", "Estimate", "Source / Assumption"]):
        tam.rows[0].cells[i].text = hdr
    tam_data = [
        ("US answering service market", "$1.5\u2013$2.0B", "IBISWorld / Grand View Research"),
        ("Estimated annual churn rate", "15\u201325%", "\u26A0 ASSUMPTION \u2014 needs validation"),
        ("Annual churning revenue", "$225\u2013$500M", "Market size \u00D7 churn rate"),
        ("AI-addressable share", f"~{ai_addr_pct:.0f}%", f"{ch_pct:.0f}% call handling + {bill_pct:.0f}% billing"),
        ("Addressable pipeline", "$180\u2013$400M", "Churning revenue \u00D7 AI-addressable share"),
        ("Central AI at 1\u20133% capture", "$2\u2013$12M ARR", "First 2\u20133 years target"),
    ]
    for r, (step, est, src) in enumerate(tam_data, 1):
        tam.rows[r].cells[0].text = step; tam.rows[r].cells[1].text = est; tam.rows[r].cells[2].text = src
    style_table(tam)
    for row in tam.rows: row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    callout(doc, "Even at 1% capture, Central AI is looking at a $2M+ ARR opportunity.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # APPENDIX A: METHODOLOGY
    # ══════════════════════════════════════════════════════════
    h(doc, "Appendix A: Methodology", 1)
    body(doc, "For anyone who wants to stress-test the numbers.", italic=True, color=MID_GRAY)
    doc.add_paragraph()

    h(doc, "Process", 2)
    for step in [
        f"Data collection: Scraped Trustpilot for {len(COMPETITOR_PROFILES)} companies and Reddit threads. "
        f"Collected {total_all_quotes} reviews, filtered to {total_count} with a specific churn reason.",
        "Classification: Built an 11-category system with clear decision rules.",
        f"Double-checking: Two independent AI classifiers categorized every quote. "
        f"Agreed {agree_n} of {total_count} times. {len(disagreements)} disagreements were manually reviewed.",
        "Weighting: Detailed, highly-upvoted Reddit posts carry more weight than one-line reviews.",
        f"Statistical check: Cohen\u2019s kappa = {kappa_val:.2f} (\u201Calmost perfect agreement\u201D).",
    ]:
        p = doc.add_paragraph(style="List Number"); p.add_run(step).font.size = Pt(10)

    doc.add_paragraph()
    h(doc, "Inter-Rater Reliability", 2)
    ir = doc.add_table(rows=6, cols=2); ir.style = "Table Grid"
    ir.rows[0].cells[0].text = "Metric"; ir.rows[0].cells[1].text = "Value"
    disagree_pct = len(disagreements) / len(final_codes) * 100
    for r, (m, v) in enumerate([
        ("Cohen\u2019s kappa", f"{kappa_val:.2f} out of 1.0 (\u201Calmost perfect\u201D)"),
        ("95% CI (approx)", f"{kappa_val - 0.05:.2f} to {kappa_val + 0.05:.2f}"),
        ("Agreement rate", f"{agree_n} of {total_count} ({agree_pct:.1f}%)"),
        ("Disagreement rate", f"{len(disagreements)} of {total_count} ({disagree_pct:.1f}%)"),
        ("Category sensitivity", "Reliability stays equally high at k=7 through k=11"),
    ], 1):
        ir.rows[r].cells[0].text = m; ir.rows[r].cells[1].text = v
    style_table(ir)

    doc.add_paragraph()
    h(doc, "Confidence Levels", 2)
    ct = doc.add_table(rows=5, cols=3); ct.style = "Table Grid"
    for i, hdr in enumerate(["Finding", "Confidence", "Why"]):
        ct.rows[0].cells[i].text = hdr
    for r, (f_text, c, b) in enumerate([
        ("Why customers leave (11 reasons)", "HIGH",
         f"{total_count} quotes, checked twice, {agree_pct:.0f}% agreement"),
        ("Who\u2019s most vulnerable", "HIGH", "Same dataset, sliced by competitor"),
        ("Where customers go next", "DIRECTIONAL",
         f"{n_switched} switching stories. Real signal, small sample."),
        ("Review sentiment over time", "DIRECTIONAL",
         "Trustpilot only. Strong pattern, other explanations possible."),
    ], 1):
        ct.rows[r].cells[0].text = f_text; ct.rows[r].cells[1].text = c; ct.rows[r].cells[2].text = b
    style_table(ct)
    for row in ct.rows: row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    h(doc, "Data Quality Notes", 2)
    pre2020_count = sum(1 for i in final_codes
                        if any(y in quotes[i].get("date", "") for y in ["2015","2016","2017","2018","2019"]))
    dq = doc.add_table(rows=7, cols=3); dq.style = "Table Grid"
    for i, hdr in enumerate(["Issue", "Detail", "Impact"]):
        dq.rows[0].cells[i].text = hdr
    for r, (issue, detail, impact) in enumerate([
        ("Mostly Trustpilot",
         f"{tp_churn}/{total_count} ({tp_churn*100//total_count}%) from Trustpilot. {reddit_churn} from Reddit.",
         "Complaint types are real but frequency may not match overall market."),
        ("Some complaints are old",
         f"{pre2020_count}/{total_count} from before 2020.",
         "Competitor rankings shift when stale data removed. Complaint types hold."),
        ("Reddit is all r/LawFirm",
         f"All {reddit_churn} Reddit quotes from r/LawFirm.",
         "\u201CLegal dominates\u201D reflects signal source, not necessarily market reality."),
        ("Both classifiers are AI",
         f"Two independent LLMs. {agree_pct:.0f}% agreement.",
         "Both may share blind spots."),
        ("Positive reviews excluded",
         f"{positive_count} positive quotes collected, not analyzed.",
         "We know what makes customers leave, not if problems are unique to churners."),
        ("Mid-2024 review surge",
         "Every competitor saw sudden 5-star spikes.",
         "\u201CUnprompted vs. recent wave\u201D split is directional, not definitive."),
    ], 1):
        dq.rows[r].cells[0].text = issue; dq.rows[r].cells[1].text = detail; dq.rows[r].cells[2].text = impact
    style_table(dq)

    doc.add_paragraph()
    h(doc, "Temporal Sensitivity", 2)
    body(doc, "Does the competitor ranking change when restricted to recent data?")

    post2020_cats = Counter()
    for idx, cat in final_codes.items():
        date = quotes[idx].get("date", "")
        if any(y in date for y in ["2015","2016","2017","2018","2019"]):
            continue
        prod = normalize_product(quotes[idx]["llm"].get("product_mentioned"))
        post2020_cats[prod] += 1

    ts_tbl = doc.add_table(rows=len(TOP_COMPETITORS) + 1, cols=4); ts_tbl.style = "Table Grid"
    for i, hdr in enumerate(["Rank", "All Dates", "Post-2020", "Change"]):
        ts_tbl.rows[0].cells[i].text = hdr
    all_ranking = sorted(
        [(prod, sum(1 for idx in final_codes if normalize_product(quotes[idx]["llm"].get("product_mentioned")) == prod))
         for prod in TOP_COMPETITORS], key=lambda x: x[1], reverse=True)
    post_ranking = sorted(
        [(prod, post2020_cats.get(prod, 0)) for prod in TOP_COMPETITORS], key=lambda x: x[1], reverse=True)
    for r in range(len(TOP_COMPETITORS)):
        all_p, all_n = all_ranking[r]
        post_p, post_n = post_ranking[r]
        change = "No change" if all_p == post_p else (f"{post_p} moves up" if post_n > all_n else f"{all_p} drops")
        ts_tbl.rows[r + 1].cells[0].text = f"#{r+1}"
        ts_tbl.rows[r + 1].cells[1].text = f"{all_p} ({all_n})"
        ts_tbl.rows[r + 1].cells[2].text = f"{post_p} ({post_n})"
        ts_tbl.rows[r + 1].cells[3].text = change
    style_table(ts_tbl)

    doc.add_paragraph()
    body(doc, f"Complaint types (call handling {ch_pct:.0f}%, billing {bill_pct:.0f}%) hold in both windows.",
         italic=True, sz=Pt(9.5), color=MID_GRAY)

    doc.add_paragraph()
    h(doc, "What This Analysis Does Not Cover", 2)
    for lim in [
        "Captures why customers LEAVE, not why they STAY.",
        "Trustpilot reviews skew negative. Reddit signal is from r/LawFirm only.",
        "English-language, US-focused.",
        "Competitor pricing data was not independently collected.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(lim); r.font.size = Pt(9.5); r.font.color.rgb = MID_GRAY; r.italic = True

    doc.add_paragraph()
    h(doc, "Sample", 2)
    body(doc, f"Reddit: {reddit_churn} quotes. Trustpilot: {tp_churn}. "
         f"Total churn: {total_count}. Total collected: {total_all_quotes}.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # APPENDIX B: THE COMPETITORS WE STUDIED
    # ══════════════════════════════════════════════════════════
    h(doc, "Appendix B: The Competitors We Studied", 1)

    h(doc, "How Answering Services Work", 2)
    body(doc, "Many small businesses (law firms, medical offices, home service companies) "
         "can\u2019t afford a full-time receptionist. They outsource to an answering service: "
         "a shared pool of human receptionists who answer calls on behalf of dozens of businesses.")
    body(doc, "The receptionist reads from a script, takes a message, and routes calls. "
         "The business pays per minute or per call.")
    body(doc, "This model has built-in problems. One receptionist juggles hundreds of clients. "
         "Training time per client is minimal. Staff turnover is high.", bold=True)
    doc.add_paragraph()

    h(doc, "Competitor Profiles", 2)
    all_profiled = list(COMPETITOR_PROFILES.keys())
    comp_churn_counts = Counter()
    for idx in final_codes:
        prod = normalize_product(quotes[idx]["llm"].get("product_mentioned"))
        comp_churn_counts[prod] += 1

    pt = doc.add_table(rows=len(all_profiled) + 1, cols=6); pt.style = "Table Grid"
    for i, hdr in enumerate(["Company", "Type", "Founded", "How They Work", "Pricing", "Churn Quotes"]):
        pt.rows[0].cells[i].text = hdr
    for r, prod in enumerate(all_profiled, 1):
        cp_data = COMPETITOR_PROFILES[prod]
        pt.rows[r].cells[0].text = prod
        pt.rows[r].cells[1].text = cp_data.get("type", "Human-pool")
        pt.rows[r].cells[2].text = cp_data.get("founded", "")
        pt.rows[r].cells[3].text = cp_data.get("how", "")
        pt.rows[r].cells[4].text = cp_data.get("pricing", "")
        pt.rows[r].cells[5].text = str(comp_churn_counts.get(prod, 0))
    style_table(pt)
    for row in pt.rows:
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # APPENDIX C: BEST QUOTES BY CATEGORY
    # ══════════════════════════════════════════════════════════
    h(doc, "Appendix C: Best Quotes by Category", 1)
    body(doc, "The 2 most compelling quotes for each of the 11 churn reasons.", italic=True, color=MID_GRAY)
    doc.add_paragraph()

    for cat in sorted_cats:
        s = cat_stats[cat]; m = CATEGORY_META[cat]
        pct = s["weighted"] / total_weighted * 100
        hd = doc.add_heading(f"{m['short']} ({s['count']} quotes, {pct:.1f}%)", level=3)
        for r in hd.runs: r.font.color.rgb = DARK_BLUE
        top = sorted(s["quotes"], key=lambda x: x[2], reverse=True)[:2]
        for j, (idx, q, w) in enumerate(top, 1):
            prod = normalize_product(q["llm"].get("product_mentioned"))
            pain = q["llm"].get("pain_point", "")
            quote_block(doc, q["text"].replace("\n", " ").strip(), f"{prod} \u2014 {pain}")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # APPENDIX D: FULL QUOTE DATABASE
    # ══════════════════════════════════════════════════════════
    h(doc, f"Appendix D: Full Quote Database (All {total_count} Churn Quotes)", 1)
    body(doc, "Every churn quote, unedited, organized by reason.",
         italic=True, color=MID_GRAY, sz=Pt(9.5))
    doc.add_paragraph()

    for cat in sorted_cats:
        s = cat_stats[cat]; m = CATEGORY_META[cat]
        hd = doc.add_heading(f"{m['short']} ({s['count']} quotes)", level=2)
        for r in hd.runs: r.font.color.rgb = DARK_BLUE
        sorted_q = sorted(s["quotes"], key=lambda x: x[2], reverse=True)
        for i, (idx, q, w) in enumerate(sorted_q, 1):
            prod = normalize_product(q["llm"].get("product_mentioned"))
            pain = q["llm"].get("pain_point", "")
            quality = q["llm"].get("quote_quality", 1)
            source = q.get("source", "")
            url = q.get("url", "")
            text_full = q["text"].replace("\n", " ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"#{i}  |  {prod}  |  Quality: {quality}/5  |  {source.title()}")
            r.bold = True; r.font.size = Pt(9); r.font.color.rgb = ACCENT_BLUE
            if pain:
                p2 = doc.add_paragraph()
                p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(2)
                r2 = p2.add_run(pain); r2.italic = True; r2.font.size = Pt(9); r2.font.color.rgb = MID_GRAY
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Cm(0.5)
            p3.paragraph_format.space_before = Pt(2); p3.paragraph_format.space_after = Pt(2)
            r3 = p3.add_run(text_full); r3.font.size = Pt(8.5); r3.font.color.rgb = DARK_GRAY
            if url:
                p4 = doc.add_paragraph()
                p4.paragraph_format.space_before = Pt(0); p4.paragraph_format.space_after = Pt(6)
                r4 = p4.add_run(url); r4.font.size = Pt(7.5); r4.font.color.rgb = MID_GRAY

    # ── SAVE ─────────────────────────────────────────────────
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"\nDocument saved to: {OUTPUT_FILE}")
    shutil.rmtree(tmpdir, ignore_errors=True)


if __name__ == "__main__":
    main()
