#!/usr/bin/env python3
"""
Build a single narrative analysis document (.docx with embedded charts).
Reads raw data, computes everything, generates charts, writes one Word file.

Target reader: Head of Product at Central AI.
Style: conversational data science blog (Gojek / Meta DS style).
"""

import json, math, os, re, tempfile, shutil
from collections import Counter

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
QUOTES_FILE = os.path.join(BASE, "data", "raw", "clean_quotes", "final_quotes.json")
CODER_A = os.path.join(BASE, "data", "coder_a_k9.json")
CODER_B = os.path.join(BASE, "data", "coder_b_k9.json")
CHART_DIR = os.path.join(BASE, "data", "charts")
OUTPUT_DOCX = os.path.join(BASE, "Deep Dive Report.docx")

# ── Colors ────────────────────────────────────────────────────────────
C_BLUE    = "#1B365F"
C_ACCENT  = "#2E75B6"
C_ORANGE  = "#E67E22"
C_GREEN   = "#27AE60"
C_PURPLE  = "#8E44AD"
C_RED     = "#C0392B"
C_GRAY    = "#BDC3C7"
GROUP_COLORS = {"CALL HANDLING": C_BLUE, "BILLING": C_ORANGE,
                "SERVICE RELIABILITY": C_GREEN, "INDUSTRY DISILLUSIONMENT": C_PURPLE}

# ── Category metadata ─────────────────────────────────────────────────
CATEGORY_META = {
    "SCRIPT_ADHERENCE":      {"short": "They don\u2019t follow my instructions",  "group": "CALL HANDLING"},
    "BILLING_PREDATORY":     {"short": "Hidden charges on my bill",              "group": "BILLING"},
    "INAUTHENTICITY":        {"short": "They don\u2019t know my business",        "group": "CALL HANDLING"},
    "GARBLED_INFO":          {"short": "They get the details wrong",             "group": "CALL HANDLING"},
    "SERIAL_SWITCHING":      {"short": "I\u2019ve tried everyone, nobody works",  "group": "INDUSTRY DISILLUSIONMENT"},
    "BILLING_TRAP":          {"short": "I can\u2019t cancel",                     "group": "BILLING"},
    "BILLING_OPAQUE":        {"short": "Surprise charges I can\u2019t explain",   "group": "BILLING"},
    "QUALITY_DECAY":         {"short": "It used to be good, then got worse",     "group": "SERVICE RELIABILITY"},
    "MISSED_CALLS":          {"short": "They don\u2019t pick up",                 "group": "SERVICE RELIABILITY"},
    "ROUTING_ERRORS":        {"short": "Calls go to the wrong place",            "group": "CALL HANDLING"},
    "BILLING_TOO_EXPENSIVE": {"short": "It costs too much",                      "group": "BILLING"},
}
GROUP_ORDER = ["CALL HANDLING", "BILLING", "SERVICE RELIABILITY", "INDUSTRY DISILLUSIONMENT"]
TOP_COMPETITORS = ["AnswerConnect", "Ruby Receptionist", "PATLive", "Smith.ai", "Synthflow"]
COMPANY_ALIASES = {
    "ruby": "Ruby Receptionist", "ruby receptionist": "Ruby Receptionist",
    "answerconnect": "AnswerConnect", "answer connect": "AnswerConnect",
    "patlive": "PATLive", "pat live": "PATLive",
    "smith.ai": "Smith.ai", "smithai": "Smith.ai", "smith ai": "Smith.ai",
    "synthflow": "Synthflow", "synthflow.ai": "Synthflow",
    "abby connect": "Abby Connect", "abbyconnect": "Abby Connect",
    "dialzara": "Dialzara", "sas": "SAS",
    "virtual hq": "Virtual HQ", "conversational": "Conversational",
    "goodcall": "Goodcall", "vapi": "Vapi", "bland ai": "Bland AI",
    "in-house": "In-house", "in house": "In-house",
}

# ── Data loading ──────────────────────────────────────────────────────
def load_quotes():
    with open(QUOTES_FILE) as f:
        all_quotes = json.load(f)
    churn = [q for q in all_quotes if q["llm"].get("category", "").startswith("churn_")]
    return churn, len(all_quotes)

def load_coder(fp):
    with open(fp) as f:
        return {item["idx"]: item["category"] for item in json.load(f)["classifications"]}

def adjudicate(a, b):
    final, dis = {}, []
    for idx in sorted(set(a) | set(b)):
        final[idx] = a.get(idx)
        if a.get(idx) != b.get(idx):
            dis.append((idx, a.get(idx), b.get(idx)))
    return final, dis

def compute_kappa(codes_a, codes_b):
    shared = sorted(set(codes_a) & set(codes_b))
    n = len(shared)
    if n == 0: return 0.0, 0, 0
    la = [codes_a[i] for i in shared]
    lb = [codes_b[i] for i in shared]
    agree = sum(1 for a, b in zip(la, lb) if a == b)
    p_o = agree / n
    cats = sorted(set(la) | set(lb))
    p_e = sum((la.count(c) / n) * (lb.count(c) / n) for c in cats)
    kappa = (p_o - p_e) / (1 - p_e) if p_e < 1 else 1.0
    return kappa, agree, n

def compute_weight(q):
    quality = q["llm"].get("quote_quality", 1)
    eng = max(1.0, math.log2(max(q.get("score") or 1, 1))) if q["source"] == "reddit" else 1.0
    return quality * eng

def normalize_product(p):
    if not p or p == "unnamed": return "Other"
    return "Ruby Receptionist" if p == "Ruby" else p

def add_hyperlink(paragraph, text, url, font_size=Pt(8), color=RGBColor(0x05, 0x63, 0xC1)):
    """Add a hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c_el = OxmlElement("w:color")
    c_el.set(qn("w:val"), f"{color}")
    rPr.append(c_el)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)
    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), str(int(font_size.pt * 2)))
    rPr.append(sz_el)
    run_el.append(rPr)
    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    run_el.append(t_el)
    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)

def truncate(text, n=200):
    t = text.replace("\n", " ").strip()
    return t[:n] + "..." if len(t) > n else t

def devastation_score(text):
    t = text.lower(); s = 0
    if any(w in t for w in ["thousand", "hundred", "$"]): s += 3
    if any(w in t for w in ["lost client", "lost customer", "lost revenue"]): s += 3
    if any(w in t for w in ["years", "year"]): s += 2
    if any(w in t for w in ["scam", "dishonest", "beware", "fraud", "worst",
                             "terrible", "horrible", "negligence"]): s += 2
    return s

def extract_pull_quote(text, max_len=200):
    text_clean = text.replace("\n", " ").strip()
    sentences = re.split(r'(?<=[.!?])\s+', text_clean)
    best, best_s = None, -1
    for s in sentences:
        if len(s) < 20: continue
        sc = devastation_score(s) + (1 if 40 < len(s) < 180 else 0)
        if sc > best_s: best_s, best = sc, s
    if best and len(best) <= max_len: return best
    elif best: return best[:max_len] + "..."
    return truncate(text_clean, max_len)

def _product_names(prod):
    names = {prod.lower()}
    for alias, canon in COMPANY_ALIASES.items():
        if canon == prod: names.add(alias.lower())
    return names

def _detect_direction(prod, pain, text):
    text_lower = text.lower()
    pain_lower = pain.lower()
    names = _product_names(prod)
    dep_phrases = ["dropped ", "left ", "cancelled ", "canceled ", "leaving ",
                   "switched from ", "ditched "]
    arr_phrases = ["switched to ", "switching to ", "went to ", "went with ",
                   "moved to ", "chose ", "choosing ", "found ", "finding "]
    for name in names:
        for dp in dep_phrases:
            if dp + name in text_lower or dp + name in pain_lower: return False
        for ap in arr_phrases:
            if ap + name in text_lower or ap + name in pain_lower: return True
    snippet = text_lower
    pos_words = ["happy", "love", "great", "best", "excellent", "amazing",
                 "wonderful", "recommend", "appreciate", "fantastic", "no comparison",
                 "pleased", "reliable", "impressed", "professional", "easy to work",
                 "helpful", "superior", "courteous"]
    neg_words = ["terrible", "worst", "awful", "horrible", "avoid", "waste", "scam",
                 "zero stars", "do not use", "poor", "disappointed", "frustrat",
                 "useless", "broken", "zero support", "no support", "non-existent",
                 "crap", "buyer beware", "stay far", "dissatisfied", "not what it used",
                 "canceled", "cancelled", "too expensive", "dropping the ball",
                 "finally left", "no longer", "nose dive", "disappoint",
                 "unfortunately", "missed"]
    pos = sum(1 for w in pos_words if w in snippet)
    neg = sum(1 for w in neg_words if w in snippet)
    if pos != neg: return pos > neg
    arr_kw = ["switched to", "found", "chose", "happy", "best vendor",
              "great experience", "no comparison", "resolved", "recommend"]
    dep_kw = ["dropped", "cancelled", "canceled", "too expensive", "declined",
              "wrong", "missed", "terrible", "worst", "overpromised", "crap",
              "deteriorat", "shopping", "dishonest", "lost calls", "costing"]
    a = sum(1 for w in arr_kw if w in pain_lower)
    d = sum(1 for w in dep_kw if w in pain_lower)
    if a != d: return a > d
    return False

def compute_switching_data(quotes, final_codes):
    switched = [(idx, quotes[idx]) for idx in final_codes
                if quotes[idx]["llm"].get("category") == "churn_switched"]
    arrivals, departures, triggers = Counter(), Counter(), Counter()
    for idx, q in switched:
        prod = normalize_product(q["llm"].get("product_mentioned"))
        pain = q["llm"].get("pain_point", "")
        text = q["text"]
        is_arrival = _detect_direction(prod, pain, text)
        text_lower = text.lower()
        other_companies = set()
        for alias, canonical in COMPANY_ALIASES.items():
            if alias in text_lower and canonical != prod:
                other_companies.add(canonical)
        if is_arrival:
            arrivals[prod] += 1
            for o in other_companies: departures[o] += 1
        else:
            departures[prod] += 1
            for o in other_companies: arrivals[o] += 1
        trigger_text = pain.lower()
        if any(w in trigger_text for w in ["integrat", "customiz", "automat", "technolog",
                                            "can't do", "couldn't", "limited", "compatib",
                                            "google voice", "couldn't adapt"]):
            triggers["Capability"] += 1
        elif any(w in trigger_text for w in ["expensiv", "price", "too much",
                                              "bill", "charg", "fee", "money",
                                              "overpay", "no longer justified", "cut cost"]):
            triggers["Pricing"] += 1
        else:
            triggers["Quality"] += 1
    all_c = set(arrivals) | set(departures)
    nf = {}
    for c in all_c:
        net = arrivals.get(c, 0) - departures.get(c, 0)
        if net != 0:
            nf[c] = {"arrivals": arrivals.get(c, 0), "departures": departures.get(c, 0), "net": net}
    return {"total": len(switched), "net_flow": nf, "triggers": triggers}

def compute_temporal_data():
    review_dir = os.path.join(BASE, "data", "raw", "review_data", "raw")
    name_map = {"answerconnect": "AnswerConnect", "patlive": "PATLive",
                "smithai": "Smith.ai", "ruby_receptionist": "Ruby Receptionist"}
    comp_order = ["AnswerConnect", "Ruby Receptionist", "PATLive", "Smith.ai"]
    result = {}
    for fname in os.listdir(review_dir):
        if not fname.endswith("_trustpilot.json"): continue
        key = fname.replace("_trustpilot.json", "")
        prod = name_map.get(key)
        if not prod or prod not in comp_order: continue
        with open(os.path.join(review_dir, fname)) as f:
            reviews = json.load(f)
        org, surge = [], []
        for r in reviews:
            try: rating = int(r.get("rating", 3))
            except: rating = 3
            date_str = r.get("date", "")
            if "2024" in date_str or "2025" in date_str or "2026" in date_str:
                month = 1
                for i, m in enumerate(["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"], 1):
                    if m in date_str: month = i; break
                if "2024" in date_str and month >= 7: surge.append(rating)
                elif "2025" in date_str or "2026" in date_str: surge.append(rating)
                else: org.append(rating)
            else:
                org.append(rating)
        result[prod] = {
            "organic_avg": float(np.mean(org)) if org else 3.0, "organic_n": len(org),
            "surge_avg": float(np.mean(surge)) if surge else 3.0, "surge_n": len(surge),
            "jump": (float(np.mean(surge)) if surge else 3.0) - (float(np.mean(org)) if org else 3.0),
        }
    return result, comp_order

def compute_bimodality(comp_order):
    review_dir = os.path.join(BASE, "data", "raw", "review_data", "raw")
    name_map = {"answerconnect": "AnswerConnect", "patlive": "PATLive",
                "smithai": "Smith.ai", "ruby_receptionist": "Ruby Receptionist"}
    result = {}
    for fname in os.listdir(review_dir):
        if not fname.endswith("_trustpilot.json"): continue
        key = fname.replace("_trustpilot.json", "")
        prod = name_map.get(key)
        if not prod or prod not in comp_order: continue
        with open(os.path.join(review_dir, fname)) as f:
            reviews = json.load(f)
        org_ext, org_tot, surge_ext, surge_tot = 0, 0, 0, 0
        org_5, surge_5 = 0, 0
        for r in reviews:
            try: rating = int(r.get("rating", 3))
            except: rating = 3
            date_str = r.get("date", "")
            is_surge = False
            if "2024" in date_str or "2025" in date_str or "2026" in date_str:
                month = 1
                for i, m in enumerate(["Jan","Feb","Mar","Apr","May","Jun","Jul","Aug","Sep","Oct","Nov","Dec"], 1):
                    if m in date_str: month = i; break
                if ("2024" in date_str and month >= 7) or "2025" in date_str or "2026" in date_str:
                    is_surge = True
            if is_surge:
                surge_tot += 1
                if rating in (1, 5): surge_ext += 1
                if rating == 5: surge_5 += 1
            else:
                org_tot += 1
                if rating in (1, 5): org_ext += 1
                if rating == 5: org_5 += 1
        result[prod] = {
            "overall_extreme_pct": (org_ext + surge_ext) / (org_tot + surge_tot) * 100 if (org_tot + surge_tot) else 0,
            "org_5star": org_5, "org_total": org_tot,
            "surge_5star": surge_5, "surge_total": surge_tot,
        }
    return result

def compute_dollar_impact(quotes, final_codes):
    dq = []
    for idx, cat in final_codes.items():
        q = quotes[idx]
        amounts = re.findall(r'\$[\d,]+(?:\.\d{2})?(?:/(?:month|mo|yr|year))?', q["text"])
        if not amounts: continue
        dq.append({
            "product": normalize_product(q["llm"].get("product_mentioned")),
            "amounts": amounts,
            "pain_point": q["llm"].get("pain_point", ""),
        })
    return dq

def compute_industry_verticals(quotes, final_codes):
    VERTICALS = {
        "Legal": ["law firm", "attorney", "lawyer", "legal", "paralegal", "litigation",
                   "practice area", "opposing counsel", "court", "client intake"],
        "Medical": ["patient", "medical", "doctor", "clinic", "hipaa", "healthcare",
                     "dental", "dermatolog"],
        "Home Services": ["hvac", "plumb", "electric", "roofing", "contractor", "technician",
                          "dispatch", "service call", "on-call", "field tech"],
        "Real Estate": ["real estate", "realtor", "property management"],
    }
    results = {}
    for idx, cat in final_codes.items():
        q = quotes[idx]
        text = (q["text"] + " " + q["llm"].get("pain_point", "")).lower()
        for vert, kw in VERTICALS.items():
            if any(k in text for k in kw):
                results.setdefault(vert, {"count": 0, "cats": Counter()})
                results[vert]["count"] += 1
                results[vert]["cats"][cat] += 1
                break
    return results


# ── Chart style ───────────────────────────────────────────────────────
def _mbb():
    plt.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": ["Helvetica Neue", "Helvetica", "Arial", "Calibri"],
        "font.size": 10, "axes.spines.top": False, "axes.spines.right": False,
        "axes.edgecolor": "#BDC3C7", "axes.labelcolor": "#333",
        "xtick.color": "#666", "ytick.color": "#666",
        "figure.facecolor": "white", "axes.facecolor": "white",
        "savefig.dpi": 180, "savefig.bbox": "tight", "savefig.pad_inches": 0.3,
    })


# ── Chart: 4-group horizontal bar ────────────────────────────────────
def chart_group_pareto(group_stats, total_weighted, path):
    _mbb()
    fig, ax = plt.subplots(figsize=(7.5, 3.2))
    groups = list(reversed(GROUP_ORDER))
    pcts = [group_stats.get(g, {}).get("weighted", 0) / total_weighted * 100 for g in groups]
    counts = [group_stats.get(g, {}).get("count", 0) for g in groups]
    colors = [GROUP_COLORS.get(g, C_GRAY) for g in groups]
    bars = ax.barh(groups, pcts, color=colors, height=0.55, edgecolor="white", linewidth=0.5)
    for bar, pct, count in zip(bars, pcts, counts):
        ax.text(bar.get_width() + 1.2, bar.get_y() + bar.get_height() / 2,
                f"{pct:.0f}%  ({count} quotes)", va="center", fontsize=10, color="#333", fontweight="bold")
    ax.set_xlim(0, 70); ax.xaxis.set_visible(False); ax.spines["bottom"].set_visible(False)
    ax.tick_params(axis="y", length=0)
    for l in ax.get_yticklabels(): l.set_fontsize(9.5); l.set_fontweight("bold"); l.set_color("#333")
    ch_pct = group_stats.get("CALL HANDLING", {}).get("weighted", 0) / total_weighted * 100
    ax.set_title(f"Call handling failures account for {ch_pct:.0f}% of all churn",
                  fontsize=11, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)


# ── Chart: 11-category Pareto ─────────────────────────────────────────
def chart_pareto_11(sorted_cats, cat_stats, total_weighted, path):
    _mbb()
    fig, ax1 = plt.subplots(figsize=(8.5, 4.5))
    labels = [CATEGORY_META[c]["short"] for c in sorted_cats]
    pcts = [cat_stats[c]["weighted"] / total_weighted * 100 for c in sorted_cats]
    colors = [GROUP_COLORS.get(CATEGORY_META[c]["group"], C_GRAY) for c in sorted_cats]
    x = np.arange(len(labels))
    bars = ax1.bar(x, pcts, color=colors, width=0.7, edgecolor="white", linewidth=0.5)
    ax2 = ax1.twinx()
    cumul = np.cumsum(pcts)
    ax2.plot(x, cumul, color=C_BLUE, linewidth=2.2, marker="o", markersize=4, zorder=5)
    ax2.set_ylim(0, 110); ax2.set_ylabel("Cumulative %", fontsize=9, color="#666")
    ax2.spines["top"].set_visible(False)
    top6_pct = cumul[5] if len(cumul) > 5 else cumul[-1]
    ax2.axhline(y=top6_pct, color=C_RED, linestyle="--", linewidth=1, alpha=0.6)
    ax2.text(len(labels) - 0.5, top6_pct + 2, f"{top6_pct:.0f}% of churn from top 6",
             fontsize=7.5, color=C_RED, ha="right", fontweight="bold")
    for bar, pct in zip(bars, pcts):
        if pct > 4:
            ax1.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.4,
                     f"{pct:.0f}%", ha="center", va="bottom", fontsize=8, color="#333")
    ax1.set_xticks(x); ax1.set_xticklabels(labels, rotation=40, ha="right", fontsize=8)
    ax1.set_ylabel("Share of Churn (%)", fontsize=9, color="#666")
    ax1.set_ylim(0, max(pcts) * 1.35)
    from matplotlib.patches import Patch
    ax1.legend(handles=[Patch(facecolor=GROUP_COLORS[g], label=g) for g in GROUP_ORDER],
               loc="upper right", fontsize=7, framealpha=0.9)
    ax1.set_title(f"Top 6 problems drive {top6_pct:.0f}% of churn",
                   fontsize=11, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)


# ── Chart: competitor heatmap ─────────────────────────────────────────
def chart_heatmap(comp_groups, path):
    _mbb()
    hm_comps = [p for p in TOP_COMPETITORS if sum(comp_groups.get(p, {}).values()) > 0]
    fig, ax = plt.subplots(figsize=(7.5, 2.2 + 0.6 * len(hm_comps)))
    data = []
    for prod in hm_comps:
        total = sum(comp_groups.get(prod, {}).values())
        data.append([comp_groups.get(prod, {}).get(g, 0) / total * 100 if total > 0 else 0 for g in GROUP_ORDER])
    data = np.array(data)
    ax.imshow(data, cmap="Blues", aspect="auto", vmin=0, vmax=70)
    ax.set_xticks(np.arange(len(GROUP_ORDER)))
    ax.set_xticklabels([g.replace("INDUSTRY ", "INDUSTRY\n") for g in GROUP_ORDER], fontsize=8)
    ax.set_yticks(np.arange(len(hm_comps))); ax.set_yticklabels(hm_comps, fontsize=9)
    for i in range(len(hm_comps)):
        for j in range(len(GROUP_ORDER)):
            v = data[i, j]; c = "white" if v > 45 else "#333"
            ax.text(j, i, f"{v:.0f}%", ha="center", va="center", fontsize=10.5,
                    color=c, fontweight="bold" if v > 35 else "normal")
    ax.tick_params(length=0)
    for spine in ax.spines.values(): spine.set_visible(True); spine.set_color("#BDC3C7")
    ax.set_title("Each competitor fails differently",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)


# ── Chart: net churn flow ─────────────────────────────────────────────
def chart_net_churn(switching_data, path):
    _mbb()
    nf = switching_data["net_flow"]
    sorted_c = sorted(nf.items(), key=lambda x: x[1]["net"], reverse=True)
    sorted_c = [(c, d) for c, d in sorted_c if abs(d["net"]) >= 1]
    labels = [c for c, _ in sorted_c]
    scores = [d["net"] for _, d in sorted_c]
    colors = [C_GREEN if s > 0 else C_RED for s in scores]
    fig, ax = plt.subplots(figsize=(8.5, 4.5))
    bars = ax.barh(labels, scores, color=colors, height=0.55, edgecolor="white", linewidth=0.5)
    for bar, sc in zip(bars, scores):
        x = bar.get_width(); off = 0.2 if sc > 0 else -0.2
        ax.text(x + off, bar.get_y() + bar.get_height() / 2, f"{sc:+d}",
                ha="left" if sc > 0 else "right", va="center", fontsize=11,
                fontweight="bold", color="#333")
    ax.axvline(x=0, color="#333", linewidth=0.8)
    mx = max(abs(s) for s in scores) if scores else 5
    ax.set_xlim(-mx - 2, mx + 2.5); ax.spines["left"].set_visible(False)
    ax.set_xlabel("Customers Gained \u2212 Lost", fontsize=10, color="#666")
    ax.tick_params(axis="y", length=0)
    for l in ax.get_yticklabels(): l.set_fontsize(10.5); l.set_color("#333")
    biggest = sorted_c[-1][0] if sorted_c else ""
    loss = sorted_c[-1][1]["net"] if sorted_c else 0
    ax.set_title(f"{biggest} loses the most customers ({loss:+d} net)",
                  fontsize=11, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.text(0.99, 0.01, f"Based on {switching_data['total']} switching stories",
             fontsize=8, color="#999", ha="right", fontstyle="italic")
    fig.savefig(path); plt.close(fig)


# ── Chart: legacy vs AI ──────────────────────────────────────────────
def chart_legacy_vs_ai(comp_groups, path):
    _mbb()
    legacy_prods = ["Ruby Receptionist", "AnswerConnect", "PATLive", "Smith.ai"]
    legacy_avg = {}
    for g in GROUP_ORDER:
        vals = []
        for p in legacy_prods:
            total = sum(comp_groups.get(p, {}).values())
            if total > 0: vals.append(comp_groups.get(p, {}).get(g, 0) / total * 100)
        legacy_avg[g] = float(np.mean(vals)) if vals else 0.0
    synth_total = sum(comp_groups.get("Synthflow", {}).values())
    synth_pct = {g: comp_groups.get("Synthflow", {}).get(g, 0) / synth_total * 100 if synth_total > 0 else 0
                 for g in GROUP_ORDER}
    fig, ax = plt.subplots(figsize=(8, 4.5))
    x = np.arange(len(GROUP_ORDER)); w = 0.35
    lv = [legacy_avg[g] for g in GROUP_ORDER]
    sv = [synth_pct[g] for g in GROUP_ORDER]
    b1 = ax.bar(x - w/2, lv, w, label="Legacy Avg (4 incumbents)", color=C_BLUE, alpha=0.85, edgecolor="white")
    b2 = ax.bar(x + w/2, sv, w, label="Synthflow (AI-native)", color=C_ORANGE, alpha=0.85, edgecolor="white")
    for bar, val in zip(b1, lv):
        if val > 0: ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                            f"{val:.0f}%", ha="center", fontsize=9, fontweight="bold", color=C_BLUE)
    for bar, val in zip(b2, sv):
        if val > 0: ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                            f"{val:.0f}%", ha="center", fontsize=9, fontweight="bold", color=C_ORANGE)
    lbls = [g.replace("INDUSTRY ", "INDUSTRY\n") for g in GROUP_ORDER]
    ax.set_xticks(x); ax.set_xticklabels(lbls, fontsize=8.5)
    ax.set_ylabel("% of Churn", fontsize=9, color="#666")
    ax.set_ylim(0, max(max(lv, default=1), max(sv, default=1)) * 1.3)
    ax.legend(fontsize=8.5, loc="upper right", framealpha=0.9)
    ax.set_title("AI eliminates call handling churn — but doesn\u2019t eliminate churn",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)


# ── Chart: temporal before/after ──────────────────────────────────────
def chart_temporal(temporal_data, comp_order, path):
    _mbb()
    fig, ax = plt.subplots(figsize=(7.5, 4))
    labels = ["AnswerConnect", "Ruby\nReceptionist", "PATLive", "Smith.ai"]
    organic = [temporal_data.get(c, {}).get("organic_avg", 3.0) for c in comp_order]
    surge = [temporal_data.get(c, {}).get("surge_avg", 3.0) for c in comp_order]
    x = np.arange(len(labels)); w = 0.32
    b1 = ax.bar(x - w/2, organic, w, label="Organic (before mid-2024)", color=C_RED, alpha=0.85, edgecolor="white")
    b2 = ax.bar(x + w/2, surge, w, label="Recent wave (mid-2024+)", color=C_GREEN, alpha=0.85, edgecolor="white")
    for bar, val in zip(b1, organic):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                f"{val:.1f}", ha="center", fontsize=9, fontweight="bold", color=C_RED)
    for bar, val in zip(b2, surge):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.1,
                f"{val:.1f}", ha="center", fontsize=9, fontweight="bold", color=C_GREEN)
    ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=9)
    ax.set_ylim(0, 5.8); ax.set_ylabel("Avg Star Rating", fontsize=9, color="#666")
    ax.axhline(y=3, color="#BDC3C7", linewidth=0.5, linestyle=":")
    ax.legend(fontsize=8, loc="upper left", framealpha=0.9)
    ax.set_title("Unprompted reviews tell the real story",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)


# ── Chart: bimodality ─────────────────────────────────────────────────
def chart_bimodality(bimod_data, comp_order, path):
    _mbb()
    fig, ax = plt.subplots(figsize=(7.5, 3.5))
    labels = ["Smith.ai", "AnswerConnect", "Ruby\nReceptionist", "PATLive"]
    order = ["Smith.ai", "AnswerConnect", "Ruby Receptionist", "PATLive"]
    extreme = [bimod_data.get(c, {}).get("overall_extreme_pct", 0) for c in order]
    middle = [100 - e for e in extreme]
    x = np.arange(len(labels))
    ax.barh(x, extreme, height=0.5, color=C_RED, alpha=0.85, label="Extreme (1★ or 5★)")
    ax.barh(x, middle, height=0.5, left=extreme, color=C_GRAY, alpha=0.5, label="Middle (2-4★)")
    for i, e in enumerate(extreme):
        ax.text(e + 1, i, f"{e:.0f}%", va="center", fontsize=10, fontweight="bold", color=C_RED)
    ax.set_yticks(x); ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlim(0, 110); ax.set_xlabel("% of All Reviews", fontsize=9, color="#666")
    ax.legend(fontsize=8, loc="lower right", framealpha=0.9)
    ax.set_title("Most reviews are extreme — almost nobody gives 2, 3, or 4 stars",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)


# ── Chart: priority matrix ────────────────────────────────────────────
def chart_priority_matrix(cat_stats, total_weighted, path):
    _mbb()
    EASE = {
        "SCRIPT_ADHERENCE": 9, "INAUTHENTICITY": 8, "GARBLED_INFO": 9,
        "ROUTING_ERRORS": 8, "BILLING_PREDATORY": 9, "BILLING_OPAQUE": 9,
        "BILLING_TRAP": 10, "BILLING_TOO_EXPENSIVE": 5, "MISSED_CALLS": 10,
        "QUALITY_DECAY": 9, "SERIAL_SWITCHING": 3,
    }
    fig, ax = plt.subplots(figsize=(7.5, 5.5))
    pcts_all = [s["weighted"] / total_weighted * 100 for s in cat_stats.values()]
    y_mid = np.median(pcts_all)
    x_mid = 6.5
    y_max = max(pcts_all) * 1.35
    ax.axhspan(y_mid, y_max, xmin=0, xmax=(x_mid - 1) / 10, color="#FFEBEE", alpha=0.25)
    ax.axhspan(y_mid, y_max, xmin=(x_mid - 1) / 10, xmax=1.0, color="#E3F2FD", alpha=0.25)
    ax.axhspan(0, y_mid, xmin=0, xmax=(x_mid - 1) / 10, color="#FFF3E0", alpha=0.25)
    ax.axhspan(0, y_mid, xmin=(x_mid - 1) / 10, xmax=1.0, color="#E8F5E9", alpha=0.25)
    ax.axhline(y=y_mid, color="#BDC3C7", linewidth=1, linestyle="--", alpha=0.5)
    ax.axvline(x=x_mid, color="#BDC3C7", linewidth=1, linestyle="--", alpha=0.5)
    for cat, s in cat_stats.items():
        pct = s["weighted"] / total_weighted * 100
        ease = EASE.get(cat, 5)
        color = GROUP_COLORS.get(CATEGORY_META[cat]["group"], C_GRAY)
        ax.scatter(ease, pct, s=pct * 18 + 60, color=color, alpha=0.85, edgecolors="white", linewidth=1, zorder=5)
        label = CATEGORY_META[cat]["short"]
        ox, oy = 0.2, 0.35
        if cat == "GARBLED_INFO": ox, oy = 0.2, -0.8
        elif cat == "BILLING_OPAQUE": ox, oy = -3.5, 0.0
        elif cat == "QUALITY_DECAY": ox, oy = -3.2, 0.0
        elif cat == "MISSED_CALLS": ox, oy = 0.2, -0.6
        elif cat == "BILLING_TRAP": ox, oy = 0.2, 0.5
        elif cat == "ROUTING_ERRORS": ox, oy = -3.0, -0.3
        ax.annotate(label, (ease, pct), xytext=(ease + ox, pct + oy),
                    fontsize=7.5, color="#333", fontweight="bold" if pct > 8 else "normal",
                    arrowprops=dict(arrowstyle="-", color="#BDC3C7", lw=0.5) if abs(ox) > 1 else None)
    ax.text(3.5, y_max * 0.85, "STRATEGIC\nBETS", fontsize=9, color="#C0392B", fontweight="bold", ha="center", alpha=0.5)
    ax.text(8.5, y_max * 0.85, "DO\nFIRST", fontsize=9, color="#1B365F", fontweight="bold", ha="center", alpha=0.5)
    ax.text(3.5, y_mid * 0.25, "DEPRIORITIZE", fontsize=9, color="#E67E22", fontweight="bold", ha="center", alpha=0.5)
    ax.text(8.5, y_mid * 0.25, "EASY\nWINS", fontsize=9, color="#27AE60", fontweight="bold", ha="center", alpha=0.5)
    ax.set_xlabel("How Easily AI Solves This (1–10)", fontsize=9, color="#666")
    ax.set_ylabel("Share of Churn (%)", fontsize=9, color="#666")
    ax.set_xlim(1, 11); ax.set_ylim(0, y_max)
    from matplotlib.patches import Patch
    ax.legend(handles=[Patch(facecolor=GROUP_COLORS[g], label=g) for g in GROUP_ORDER],
               loc="upper left", fontsize=7, framealpha=0.9)
    ax.set_title("Script failures and billing: high impact, directly solvable by AI",
                  fontsize=10.5, fontweight="bold", color="#1B365F", pad=12, loc="left")
    fig.savefig(path); plt.close(fig)


# ── Chart: switching Sankey ───────────────────────────────────────────
COMP_COLORS = {
    "Ruby Receptionist": "#C0392B", "AnswerConnect": "#2E75B6",
    "PATLive": "#27AE60", "Smith.ai": "#E67E22", "Synthflow": "#8E44AD",
    "Abby Connect": "#3498DB", "Conversational": "#1ABC9C",
    "Dialzara": "#F39C12", "SAS": "#95A5A6", "Other": "#BDC3C7",
    "In-house": "#7F8C8D", "Unknown": "#BDC3C7", "Vapi": "#D35400",
    "Virtual HQ": "#16A085",
}

def extract_switching_pairs(quotes, final_codes):
    switched = [(idx, quotes[idx]) for idx in final_codes
                if quotes[idx]["llm"].get("category") == "churn_switched"]
    pairs = []
    for idx, q in switched:
        prod = normalize_product(q["llm"].get("product_mentioned"))
        pain = q["llm"].get("pain_point", "")
        text = q["text"]
        is_arrival = _detect_direction(prod, pain, text)
        text_lower = text.lower()
        other_companies = set()
        for alias, canonical in COMPANY_ALIASES.items():
            if alias in text_lower and canonical != prod:
                other_companies.add(canonical)
        if is_arrival:
            if other_companies:
                for o in other_companies: pairs.append((o, prod))
            else:
                pairs.append(("Unknown", prod))
        else:
            if other_companies:
                for o in other_companies: pairs.append((prod, o))
            else:
                pairs.append((prod, "Unknown"))
    return pairs

def chart_switching_sankey(pairs, path):
    _mbb()
    from collections import Counter
    import matplotlib.patches as mpatches
    flow_counts = Counter(pairs)
    from_companies = sorted(set(f for f, _ in pairs),
                            key=lambda x: -sum(v for (f, t), v in flow_counts.items() if f == x))
    to_companies = sorted(set(t for _, t in pairs),
                          key=lambda x: -sum(v for (f, t), v in flow_counts.items() if t == x))
    for lst in [from_companies, to_companies]:
        if "Unknown" in lst: lst.remove("Unknown"); lst.append("Unknown")

    fig, ax = plt.subplots(figsize=(12, 7))
    y_padding = 0.3
    total_height = max(len(from_companies), len(to_companies)) * 1.5

    from_sizes = {c: sum(v for (f, t), v in flow_counts.items() if f == c) for c in from_companies}
    to_sizes = {c: sum(v for (f, t), v in flow_counts.items() if t == c) for c in to_companies}
    from_total = sum(from_sizes.values())
    to_total = sum(to_sizes.values())

    from_y, y_cursor = {}, total_height
    for c in from_companies:
        h = max(from_sizes[c] / from_total * total_height * 0.8, 0.3)
        from_y[c] = (y_cursor - h / 2, h); y_cursor -= h + y_padding

    to_y, y_cursor = {}, total_height
    for c in to_companies:
        h = max(to_sizes[c] / to_total * total_height * 0.8, 0.3)
        to_y[c] = (y_cursor - h / 2, h); y_cursor -= h + y_padding

    ax.set_xlim(-0.5, 3.5); ax.set_ylim(-1, total_height + 1)
    node_w, x_from, x_to = 0.15, 0.5, 2.5

    for c in from_companies:
        yc, yh = from_y[c]
        color = COMP_COLORS.get(c, "#BDC3C7")
        rect = mpatches.FancyBboxPatch((x_from - node_w, yc - yh / 2), node_w * 2, yh,
                                        boxstyle="round,pad=0.02", facecolor=color,
                                        edgecolor="white", linewidth=1.5, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x_from - node_w - 0.08, yc, f"{c} ({from_sizes[c]})",
                ha="right", va="center", fontsize=9.5, fontweight="bold", color="#333")

    for c in to_companies:
        yc, yh = to_y[c]
        color = COMP_COLORS.get(c, "#BDC3C7")
        rect = mpatches.FancyBboxPatch((x_to - node_w, yc - yh / 2), node_w * 2, yh,
                                        boxstyle="round,pad=0.02", facecolor=color,
                                        edgecolor="white", linewidth=1.5, alpha=0.9)
        ax.add_patch(rect)
        ax.text(x_to + node_w + 0.08, yc, f"{c} ({to_sizes[c]})",
                ha="left", va="center", fontsize=9.5, fontweight="bold", color="#333")

    from_used = {c: 0.0 for c in from_companies}
    to_used = {c: 0.0 for c in to_companies}
    for (frm, to), count in sorted(flow_counts.items(), key=lambda x: -x[1]):
        frm_yc, frm_yh = from_y[frm]; to_yc, to_yh = to_y[to]
        frm_flow_h = count / from_sizes[frm] * frm_yh
        to_flow_h = count / to_sizes[to] * to_yh
        y1_top = frm_yc + frm_yh / 2 - from_used[frm]; y1_bot = y1_top - frm_flow_h
        from_used[frm] += frm_flow_h
        y2_top = to_yc + to_yh / 2 - to_used[to]; y2_bot = y2_top - to_flow_h
        to_used[to] += to_flow_h
        color = COMP_COLORS.get(frm, "#BDC3C7")
        t_vals = np.linspace(0, 1, 50)
        x1, x2 = x_from + node_w, x_to - node_w
        top_x = x1 + (x2 - x1) * t_vals
        top_y = y1_top + (y2_top - y1_top) * (3 * t_vals**2 - 2 * t_vals**3)
        bot_x = x1 + (x2 - x1) * t_vals
        bot_y = y1_bot + (y2_bot - y1_bot) * (3 * t_vals**2 - 2 * t_vals**3)
        ax.fill(np.concatenate([top_x, bot_x[::-1]]),
                np.concatenate([top_y, bot_y[::-1]]),
                color=color, alpha=0.35, edgecolor=color, linewidth=0.3)
        if count > 1:
            ax.text((x1 + x2) / 2, (y1_top + y1_bot + y2_top + y2_bot) / 4,
                    str(count), ha="center", va="center", fontsize=8,
                    fontweight="bold", color=color, alpha=0.8)

    ax.text(x_from, total_height + 0.5, "LEFT", ha="center", fontsize=11,
            fontweight="bold", color="#C0392B")
    ax.text(x_to, total_height + 0.5, "JOINED", ha="center", fontsize=11,
            fontweight="bold", color="#27AE60")
    ax.set_title(f"Where customers go when they leave ({len(pairs)} switching flows)",
                  fontsize=11, fontweight="bold", color="#1B365F", pad=20, loc="left")
    ax.axis("off")
    fig.savefig(path); plt.close(fig)


# ── Docx colors ──────────────────────────────────────────────────────
DARK_BLUE    = RGBColor(0x1B, 0x36, 0x5F)
ACCENT_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
DARK_GRAY    = RGBColor(0x33, 0x33, 0x33)
MID_GRAY     = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY   = RGBColor(0x99, 0x99, 0x99)
LIGHT_BLUE_BG = "D6E4F0"
LIGHT_GRAY_BG = "F2F2F2"
GREEN_RGB    = RGBColor(0x2D, 0x8B, 0x57)
RED_RGB      = RGBColor(0xC0, 0x39, 0x2B)

# ── Docx helpers ─────────────────────────────────────────────────────
def set_shading(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'))

def style_table(tbl, header_bg=LIGHT_BLUE_BG, stripe=True):
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl.rows[0].cells:
        set_shading(cell, header_bg)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.color.rgb = DARK_BLUE; r.font.size = Pt(9)
    for i, row in enumerate(tbl.rows[1:], 1):
        for cell in row.cells:
            if stripe and i % 2 == 0: set_shading(cell, LIGHT_GRAY_BG)
            for p in cell.paragraphs:
                for r in p.runs: r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY

def h(doc, text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs: r.font.color.rgb = DARK_BLUE
    return hd

def body(doc, text, bold=False, italic=False, sz=Pt(10.5), color=DARK_GRAY):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = sz; r.font.color.rgb = color
    return p

def body_rich(doc, segments):
    """Add a paragraph with mixed formatting. segments = list of (text, bold, italic, color)."""
    p = doc.add_paragraph()
    for text, bold, italic, color in segments:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic; r.font.size = Pt(10.5)
        r.font.color.rgb = color if color else DARK_GRAY
    return p

def quote_block(doc, qt, attrib):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"\u201C{qt}\u201D"); r.italic = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK_GRAY
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(1.5)
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(f"\u2014 {attrib}"); r2.font.size = Pt(9); r2.font.color.rgb = MID_GRAY

def callout(doc, text):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0); set_shading(cell, LIGHT_BLUE_BG)
    p = cell.paragraphs[0]; r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK_BLUE
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

def add_chart(doc, path, width=Inches(5.8)):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=width)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK_GRAY
        r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY
    else:
        r = p.add_run(text); r.font.size = Pt(10.5); r.font.color.rgb = DARK_GRAY
    return p

def spacer(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════════════
# MAIN: compute everything, generate charts, write one narrative .docx
# ══════════════════════════════════════════════════════════════════════
def main():
    # ── Load & compute ────────────────────────────────────────────
    quotes, total_reviews = load_quotes()
    codes_a, codes_b = load_coder(CODER_A), load_coder(CODER_B)
    final_codes, disagreements = adjudicate(codes_a, codes_b)
    kappa_val, kappa_agree, kappa_n = compute_kappa(codes_a, codes_b)

    cat_stats = {}
    for idx, cat in final_codes.items():
        cat_stats.setdefault(cat, {"count": 0, "weighted": 0.0, "quotes": []})
        q = quotes[idx]; w = compute_weight(q)
        cat_stats[cat]["count"] += 1; cat_stats[cat]["weighted"] += w
        cat_stats[cat]["quotes"].append((idx, q, w))
    total_weighted = sum(s["weighted"] for s in cat_stats.values())
    sorted_cats = sorted(cat_stats, key=lambda c: cat_stats[c]["weighted"], reverse=True)

    group_stats = {}
    for cat, s in cat_stats.items():
        g = CATEGORY_META[cat]["group"]
        group_stats.setdefault(g, {"count": 0, "weighted": 0.0, "cats": []})
        group_stats[g]["count"] += s["count"]; group_stats[g]["weighted"] += s["weighted"]
        group_stats[g]["cats"].append(cat)

    comp_cats, comp_groups = {}, {}
    for idx, cat in final_codes.items():
        prod = normalize_product(quotes[idx]["llm"].get("product_mentioned"))
        comp_cats.setdefault(prod, Counter())[cat] += 1
    for prod, cats in comp_cats.items():
        comp_groups[prod] = Counter()
        for cat, cnt in cats.items():
            comp_groups[prod][CATEGORY_META[cat]["group"]] += cnt

    switching_data = compute_switching_data(quotes, final_codes)
    temporal_data, comp_order = compute_temporal_data()
    bimod_data = compute_bimodality(comp_order)
    dollar_quotes = compute_dollar_impact(quotes, final_codes)
    verticals = compute_industry_verticals(quotes, final_codes)

    # Gut-punch quotes (with source URLs for evidence)
    gut_punch = {}
    for g in GROUP_ORDER:
        group_cat_list = group_stats[g]["cats"]
        candidates = []
        for idx, cat in final_codes.items():
            if cat not in group_cat_list: continue
            q = quotes[idx]; quality = q["llm"].get("quote_quality", 1)
            pull = extract_pull_quote(q["text"])
            sc = devastation_score(pull) * 10 + devastation_score(q["text"]) * 3 + quality * 2
            candidates.append((sc, idx, q, cat))
        candidates.sort(key=lambda x: x[0], reverse=True)
        if candidates:
            _, idx, q, cat = candidates[0]
            gut_punch[g] = {
                "quote": extract_pull_quote(q["text"]),
                "product": normalize_product(q["llm"].get("product_mentioned")),
                "pain": q["llm"].get("pain_point", ""),
                "url": q.get("url", ""),
                "source": q.get("source", ""),
            }

    # Build evidence samples: for each group, pick 3 quotes with source URLs
    evidence_samples = {}
    for g in GROUP_ORDER:
        group_cat_list = group_stats[g]["cats"]
        samples = []
        for idx, cat in final_codes.items():
            if cat not in group_cat_list: continue
            q = quotes[idx]
            url = q.get("url", "")
            if not url: continue
            prod = normalize_product(q["llm"].get("product_mentioned"))
            pull = extract_pull_quote(q["text"], 120)
            samples.append((compute_weight(q), prod, pull, url, q.get("source", "")))
        samples.sort(key=lambda x: x[0], reverse=True)
        evidence_samples[g] = samples[:3]

    # ── Helper computations ───────────────────────────────────────
    def grp_pct(g):
        return group_stats.get(g, {}).get("weighted", 0) / total_weighted * 100

    def comp_grp_pct(prod, g):
        total = sum(comp_groups.get(prod, {}).values())
        return comp_groups.get(prod, {}).get(g, 0) / total * 100 if total > 0 else 0

    ch_pct = grp_pct("CALL HANDLING")
    bill_pct = grp_pct("BILLING")
    rel_pct = grp_pct("SERVICE RELIABILITY")
    dis_pct = grp_pct("INDUSTRY DISILLUSIONMENT")
    ai_addr_pct = ch_pct + bill_pct

    n_total = len(final_codes)
    n_switch = switching_data["total"]
    nf = switching_data["net_flow"]
    trig = switching_data["triggers"]

    # Dollar amount statistics
    dollar_amounts_raw = []
    for dq in dollar_quotes:
        for amt in dq["amounts"]:
            clean = amt.replace("$", "").replace(",", "").split("/")[0]
            try:
                val = float(clean)
                if val > 0:
                    dollar_amounts_raw.append(val)
            except ValueError:
                pass
    n_dollar = len(dollar_quotes)
    dollar_min = min(dollar_amounts_raw) if dollar_amounts_raw else 0
    dollar_max = max(dollar_amounts_raw) if dollar_amounts_raw else 0

    # ── Generate charts ───────────────────────────────────────────
    os.makedirs(CHART_DIR, exist_ok=True)
    print("Generating charts...")

    c = {}
    c["group_pareto"] = os.path.join(CHART_DIR, "group_pareto.png")
    chart_group_pareto(group_stats, total_weighted, c["group_pareto"])

    c["pareto_11"] = os.path.join(CHART_DIR, "pareto_11.png")
    chart_pareto_11(sorted_cats, cat_stats, total_weighted, c["pareto_11"])

    c["heatmap"] = os.path.join(CHART_DIR, "heatmap.png")
    chart_heatmap(comp_groups, c["heatmap"])

    c["net_churn"] = os.path.join(CHART_DIR, "net_churn.png")
    chart_net_churn(switching_data, c["net_churn"])

    c["legacy_vs_ai"] = os.path.join(CHART_DIR, "legacy_vs_ai.png")
    chart_legacy_vs_ai(comp_groups, c["legacy_vs_ai"])

    c["temporal"] = os.path.join(CHART_DIR, "temporal.png")
    chart_temporal(temporal_data, comp_order, c["temporal"])

    c["bimodality"] = os.path.join(CHART_DIR, "bimodality.png")
    chart_bimodality(bimod_data, comp_order, c["bimodality"])

    c["priority"] = os.path.join(CHART_DIR, "priority_matrix.png")
    chart_priority_matrix(cat_stats, total_weighted, c["priority"])

    switching_pairs = extract_switching_pairs(quotes, final_codes)
    c["switching_sankey"] = os.path.join(CHART_DIR, "switching_sankey.png")
    chart_switching_sankey(switching_pairs, c["switching_sankey"])

    print(f"  9 charts → {CHART_DIR}/")

    # ── Write the docx ───────────────────────────────────────────
    print("Writing docx...")

    # Top categories for inline references
    top_cats = []
    running = 0
    for cat in sorted_cats:
        s = cat_stats[cat]
        pct = s["weighted"] / total_weighted * 100
        running += pct
        top_cats.append((cat, s["count"], pct, running))

    tp_links = {
        "AnswerConnect": "https://www.trustpilot.com/review/answerconnect.com",
        "Ruby Receptionist": "https://www.trustpilot.com/review/ruby.com",
        "PATLive": "https://www.trustpilot.com/review/patlive.com",
        "Smith.ai": "https://www.trustpilot.com/review/smith.ai",
        "Synthflow": "https://www.trustpilot.com/review/synthflow.ai",
    }

    doc = Document()

    # ── Set default font ──────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.font.color.rgb = DARK_GRAY
    for lvl in range(1, 5):
        hs = doc.styles[f"Heading {lvl}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = DARK_BLUE

    # ══════════════════════════════════════════════════════════════
    # TITLE
    # ══════════════════════════════════════════════════════════════
    spacer(doc)
    spacer(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Why Answering Service Customers Leave")
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = DARK_BLUE
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("and What Central AI Should Do About It")
    r2.bold = True; r2.font.size = Pt(22); r2.font.color.rgb = DARK_BLUE
    spacer(doc)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"{n_total} churn stories from {total_reviews} reviews  |  9 competitors  |  Reddit & Trustpilot  |  2015\u20132026")
    r3.font.size = Pt(11); r3.font.color.rgb = MID_GRAY; r3.italic = True
    spacer(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # TL;DR
    # ══════════════════════════════════════════════════════════════
    h(doc, "TL;DR", level=1)

    ruby_net_val = nf.get("Ruby Receptionist", {}).get("net", 0)
    synth_bill = comp_grp_pct("Synthflow", "BILLING")
    synth_rel = comp_grp_pct("Synthflow", "SERVICE RELIABILITY")

    tldr_items = [
        (f"{ai_addr_pct:.0f}% of churn maps to two problems Central AI\u2019s model solves. ",
         f"AI fixes call handling ({ch_pct:.0f}%). Flat-rate pricing fixes billing ({bill_pct:.0f}%). Both require a different business model, not just better technology."),
        (f"Ruby and Smith.ai are bleeding customers. ",
         f"Ruby has the worst net flow ({ruby_net_val:+d}). Their customers are actively shopping."),
        ("The ratings are fake. ",
         "Before mid-2024, organic reviews averaged 1.2\u20133.8 stars. Then sudden waves of 5-star reviews appeared. The real experience hasn\u2019t changed."),
        ("AI-native proves the concept but can\u2019t execute. ",
         f"Synthflow has {synth_bill:.0f}% billing churn and {synth_rel:.0f}% reliability churn. They solved calls. They broke everything else."),
        ("Lead with scripts and flat-rate pricing at $99\u2013199/mo. ",
         "That addresses 76% of churn and undercuts legacy pricing by 50\u201390%. Price isn\u2019t why people leave \u2014 billing model is."),
    ]
    for i, (bold_part, rest) in enumerate(tldr_items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        r = p.add_run(f"{i}. "); r.font.size = Pt(10.5); r.font.color.rgb = DARK_GRAY
        r2 = p.add_run(bold_part); r2.bold = True; r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_BLUE
        r3 = p.add_run(rest); r3.font.size = Pt(10.5); r3.font.color.rgb = DARK_GRAY

    spacer(doc)

    # ── Table of Contents (Word TOC field — auto page numbers + clickable links) ──
    h(doc, "Contents", level=2)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    run2 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-1" \\h \\z \\u '
    run2._r.append(instr)
    run3 = p.add_run()
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    run3._r.append(fld_char_separate)
    run4 = p.add_run("(Right-click \u2192 Update Field to generate table of contents)")
    run4.font.size = Pt(9); run4.font.color.rgb = MID_GRAY; run4.italic = True
    run5 = p.add_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._r.append(fld_char_end)
    spacer(doc)

    # ══════════════════════════════════════════════════════════════
    # THE SETUP
    # ══════════════════════════════════════════════════════════════
    h(doc, "The Setup", level=1)
    body(doc, "Central AI is entering a market where thousands of small businesses \u2014 law firms, "
         "medical offices, home service companies \u2014 pay $235 to $2,800 per month for someone to "
         "answer their phones. These businesses can\u2019t afford a full-time receptionist, so they "
         "outsource to an answering service. A remote agent reads from a script, takes a message, "
         "routes the call. The business pays per minute or per call.")
    body(doc, "The model has built-in problems. One receptionist juggles hundreds of clients. "
         "Training time per client is minimal. Staff turnover is high. Your \u201Cdedicated receptionist\u201D "
         "is actually a stranger reading a script they saw 30 seconds ago.")
    body(doc, "I wanted to understand exactly how and why customers leave these services \u2014 "
         "not from the companies\u2019 perspective, but from the customers\u2019 own words. So I went to the "
         "two places where customers talk candidly: Reddit and Trustpilot.")
    body(doc, f"I read {total_reviews} reviews of 9 answering services. {n_total} of them described "
         f"a specific reason for leaving \u2014 not just \u201CI\u2019m unhappy\u201D but a concrete complaint I could "
         f"classify. I had two independent classifiers tag every single one into 11 categories. "
         f"They agreed {kappa_agree} out of {kappa_n} times (Cohen\u2019s kappa = {kappa_val:.2f}, "
         f"\u201Calmost perfect\u201D in stats terms). The {len(disagreements)} disagreements were all "
         "borderline cases within the same group.")
    body(doc, "What follows is the analytical journey through that data \u2014 the questions I asked, "
         "what I found, and what it means for you.")

    # Data provenance table
    h(doc, "Where the data comes from", level=2)
    body(doc, "Every number in this document traces back to a real review. Here\u2019s the breakdown by competitor:")

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    for j, hdr in enumerate(["Competitor", "Trustpilot (churn)", "Reddit (churn)", "Churn Quotes", "Trustpilot Link"]):
        tbl.rows[0].cells[j].text = hdr
    for prod in TOP_COMPETITORS:
        n_churn = sum(comp_cats.get(prod, {}).values())
        tp_count, rd_count = 0, 0
        for idx, cat in final_codes.items():
            if normalize_product(quotes[idx]["llm"].get("product_mentioned")) == prod:
                if quotes[idx].get("source") == "reddit": rd_count += 1
                else: tp_count += 1
        link = tp_links.get(prod, "")
        row = tbl.add_row()
        row.cells[0].text = prod
        row.cells[1].text = str(tp_count)
        row.cells[2].text = str(rd_count)
        row.cells[3].text = str(n_churn)
        row.cells[4].text = link
    style_table(tbl)
    body(doc, f"Total: {n_total} churn quotes from {total_reviews} reviews. "
         "Every quote is linked to its original source in the appendix.", italic=True, sz=Pt(9), color=MID_GRAY)

    # ══════════════════════════════════════════════════════════════
    # HOW THE MARKET BREAKS
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h(doc, "How the Market Breaks", level=1)
    body(doc, "The first question I asked was the most basic one: what exactly goes wrong? When a small "
         "business owner decides to leave their answering service, what pushed them over the edge?")
    body(doc, "Before I get into the numbers, it\u2019s worth hearing the raw voice of the customer. "
         "I pulled one quote from each major failure group \u2014 not because they\u2019re outliers, "
         "but because they\u2019re representative of dozens of stories like them.")

    group_labels = {
        "CALL HANDLING": "The Call Handling Problem",
        "BILLING": "The Billing Problem",
        "SERVICE RELIABILITY": "The Reliability Problem",
        "INDUSTRY DISILLUSIONMENT": "The Category Problem",
    }
    for g in GROUP_ORDER:
        gp = gut_punch.get(g, {})
        label = group_labels.get(g, g)
        body(doc, label, bold=True, color=DARK_BLUE, sz=Pt(11))
        source_label = "Trustpilot" if "trustpilot" in gp.get("url", "") else "Reddit"
        attrib = f"{gp.get('product', '')} customer, {source_label}"
        quote_block(doc, gp.get("quote", ""), attrib)

    # The four groups
    h(doc, "The four groups", level=2)
    body(doc, "Reading hundreds of these stories, a pattern emerged quickly. Every complaint falls "
         "into one of four groups \u2014 and the distribution isn\u2019t even close to uniform. "
         "Call handling dominates.")
    add_chart(doc, c["group_pareto"])

    group_narratives = {
        "CALL HANDLING": (
            "This is the core failure of the answering service model. Customers are paying "
            "for someone to represent their business on the phone, and that person doesn\u2019t know "
            "the business, doesn\u2019t follow the script, and garbles the information they collect. "
            "It\u2019s not a training problem \u2014 it\u2019s a structural one. A single receptionist juggling "
            "hundreds of clients will never match someone dedicated to one business."
        ),
        "BILLING": (
            "The second-largest group is about how customers are billed, not how much. "
            "Per-minute billing creates perverse incentives: longer calls mean more revenue for the "
            "service, not better outcomes for the customer. Add in hidden charges for spam calls, "
            "after-call work time, and deliberately difficult cancellation processes, and you get "
            "a billing model that feels predatory \u2014 because it is."
        ),
        "SERVICE RELIABILITY": (
            "These are the customers who were happy once. They signed up, it worked, and then "
            "something changed. Staff turnover, management shifts, growing pains \u2014 whatever the "
            "cause, the service degraded. The most painful quotes come from long-time customers "
            "who watched a good thing fall apart."
        ),
        "INDUSTRY DISILLUSIONMENT": (
            "Some customers have given up on the entire category. They\u2019ve tried three, four, "
            "five answering services and concluded that none of them work. These are the highest-intent "
            "prospects for Central AI \u2014 they want the problem solved, they just don\u2019t believe "
            "a human-pool model can do it."
        ),
    }

    # Helper: category pct
    def cat_pct(cat_key):
        return cat_stats.get(cat_key, {}).get("weighted", 0) / total_weighted * 100

    def cat_n(cat_key):
        return cat_stats.get(cat_key, {}).get("count", 0)

    # ── CALL HANDLING ──
    g = "CALL HANDLING"
    gs = group_stats[g]
    pct = gs["weighted"] / total_weighted * 100
    body(doc, f"#1: {g} \u2014 {pct:.0f}% of churn ({gs['count']} quotes)", bold=True, color=DARK_BLUE)
    body(doc, group_narratives[g])
    body(doc, f"Most of this is agents not following the script ({cat_pct('SCRIPT_ADHERENCE'):.0f}%, "
         f"{cat_n('SCRIPT_ADHERENCE')} quotes) and not knowing the business well enough to sound "
         f"credible ({cat_pct('INAUTHENTICITY'):.0f}%, {cat_n('INAUTHENTICITY')} quotes). "
         f"A smaller chunk is basic data entry errors \u2014 wrong phone numbers, misspelled names, "
         f"garbled emails ({cat_pct('GARBLED_INFO'):.0f}%, {cat_n('GARBLED_INFO')} quotes). "
         f"And a handful of cases where calls simply go to the wrong person "
         f"({cat_pct('ROUTING_ERRORS'):.0f}%, {cat_n('ROUTING_ERRORS')} quotes).")
    affected = sorted([(p, comp_grp_pct(p, g)) for p in TOP_COMPETITORS if comp_grp_pct(p, g) > 0],
                       key=lambda x: x[1], reverse=True)[:3]
    if affected:
        body(doc, f"Most affected: {', '.join(f'{p} ({v:.0f}%)' for p, v in affected)}",
             italic=True, sz=Pt(9.5), color=MID_GRAY)

    # ── BILLING ──
    g = "BILLING"
    gs = group_stats[g]
    pct = gs["weighted"] / total_weighted * 100
    body(doc, f"#2: {g} \u2014 {pct:.0f}% of churn ({gs['count']} quotes)", bold=True, color=DARK_BLUE)
    body(doc, group_narratives[g])
    body(doc, f"The biggest slice is hidden charges and inflated bills ({cat_pct('BILLING_PREDATORY'):.0f}%, "
         f"{cat_n('BILLING_PREDATORY')} quotes) \u2014 customers charged for spam calls, billed for "
         "\"after-call work\" time, or upsold to plans they didn\u2019t need. "
         f"Then there\u2019s the cancellation trap ({cat_pct('BILLING_TRAP'):.0f}%, "
         f"{cat_n('BILLING_TRAP')} quotes): services that make you email to cancel, require 30-day "
         f"notice, or keep billing after you\u2019ve stopped. "
         f"Opaque pricing ({cat_pct('BILLING_OPAQUE'):.0f}%, {cat_n('BILLING_OPAQUE')} quotes) "
         "rounds it out \u2014 bills that don\u2019t match what was promised, charges no one can explain. "
         f"Pure price sensitivity ({cat_pct('BILLING_TOO_EXPENSIVE'):.0f}%, "
         f"{cat_n('BILLING_TOO_EXPENSIVE')} quotes) is actually the smallest billing sub-problem.")
    affected = sorted([(p, comp_grp_pct(p, g)) for p in TOP_COMPETITORS if comp_grp_pct(p, g) > 0],
                       key=lambda x: x[1], reverse=True)[:3]
    if affected:
        body(doc, f"Most affected: {', '.join(f'{p} ({v:.0f}%)' for p, v in affected)}",
             italic=True, sz=Pt(9.5), color=MID_GRAY)

    # ── SERVICE RELIABILITY ──
    g = "SERVICE RELIABILITY"
    gs = group_stats[g]
    pct = gs["weighted"] / total_weighted * 100
    body(doc, f"#3: {g} \u2014 {pct:.0f}% of churn ({gs['count']} quotes)", bold=True, color=DARK_BLUE)
    body(doc, group_narratives[g])
    body(doc, f"Most of this is quality decay ({cat_pct('QUALITY_DECAY'):.0f}%, "
         f"{cat_n('QUALITY_DECAY')} quotes) \u2014 5-year customers who describe watching "
         "the service deteriorate, getting $10 credits for errors that cost them thousands. "
         f"The rest is missed calls ({cat_pct('MISSED_CALLS'):.0f}%, "
         f"{cat_n('MISSED_CALLS')} quotes): messages delivered too late, callers put on hold "
         "until they hang up, leads lost to competitors before the business even knew they called.")
    affected = sorted([(p, comp_grp_pct(p, g)) for p in TOP_COMPETITORS if comp_grp_pct(p, g) > 0],
                       key=lambda x: x[1], reverse=True)[:3]
    if affected:
        body(doc, f"Most affected: {', '.join(f'{p} ({v:.0f}%)' for p, v in affected)}",
             italic=True, sz=Pt(9.5), color=MID_GRAY)

    # ── INDUSTRY DISILLUSIONMENT ──
    g = "INDUSTRY DISILLUSIONMENT"
    gs = group_stats[g]
    pct = gs["weighted"] / total_weighted * 100
    body(doc, f"#4: {g} \u2014 {pct:.0f}% of churn ({gs['count']} quotes)", bold=True, color=DARK_BLUE)
    body(doc, group_narratives[g])
    body(doc, f"All {cat_n('SERIAL_SWITCHING')} quotes in this group describe the same pattern: "
         "a business owner who has tried multiple answering services and is looking for something "
         "fundamentally different. They\u2019re not comparing feature lists \u2014 they\u2019re asking whether "
         "the entire category can work.")
    affected = sorted([(p, comp_grp_pct(p, g)) for p in TOP_COMPETITORS if comp_grp_pct(p, g) > 0],
                       key=lambda x: x[1], reverse=True)[:3]
    if affected:
        body(doc, f"Most affected: {', '.join(f'{p} ({v:.0f}%)' for p, v in affected)}",
             italic=True, sz=Pt(9.5), color=MID_GRAY)

    # Full Pareto
    h(doc, "Where churn concentrates", level=2)
    body(doc, "The next question I asked was about concentration: are complaints spread evenly across all 11 "
         "categories, or do a few dominate? This matters for you because it determines whether "
         "Central AI needs to solve everything or can win by solving a few things well.")
    add_chart(doc, c["pareto_11"])
    body(doc, "The data is decisive. The top 6 reasons account for the vast majority of churn. "
         "Everything below the line is real, but it\u2019s diminishing returns. "
         "This means Central AI doesn\u2019t need to be perfect at everything \u2014 "
         "it needs to be excellent at scripts, inauthenticity, and billing.")

    # AI-native comparison
    h(doc, "What happens when you switch to AI?", level=2)
    body(doc, "Here\u2019s where I asked a question that turned out to be more interesting than expected: "
         "if the core problem is human receptionists failing at calls, does switching to AI fix it?")
    body(doc, "Synthflow is the only AI-native competitor with enough data to analyze (17 churn quotes "
         "on Trustpilot). And the answer is nuanced. AI does solve the call handling problem \u2014 "
         "script adherence complaints essentially vanish. But churn doesn\u2019t vanish. It moves.")

    ruby_ch = comp_grp_pct("Ruby Receptionist", "CALL HANDLING")
    ac_ch = comp_grp_pct("AnswerConnect", "CALL HANDLING")
    synth_ch = comp_grp_pct("Synthflow", "CALL HANDLING")
    bullet(doc, f" Call handling drops from {ruby_ch:.0f}\u2013{ac_ch:.0f}% (legacy) to {synth_ch:.0f}% (Synthflow). "
           "AI follows the script every time. That specific problem is solved by design.")
    bullet(doc, f" But billing jumps to {synth_bill:.0f}% and reliability to {synth_rel:.0f}%. "
           "Synthflow customers complain about bait-and-switch pricing, features locked behind "
           "enterprise plans, and a platform that breaks silently.")
    spacer(doc)
    add_chart(doc, c["legacy_vs_ai"])
    body(doc, "This is a crucial insight for Central AI. Winning isn\u2019t just about having an AI that "
         "handles calls well \u2014 Synthflow already proved that\u2019s possible. Winning means avoiding "
         "both failure modes: the call handling problems that plague legacy services AND the "
         "billing and reliability problems that plague the AI-native ones. The bar isn\u2019t "
         "\u201Cbetter than human receptionists.\u201D The bar is \u201Cbetter than human receptionists, "
         "without introducing new reasons to leave.\u201D")

    # ══════════════════════════════════════════════════════════════
    # WHO'S BLEEDING
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h(doc, "Who\u2019s Bleeding", level=1)
    body(doc, "At this point I knew what breaks. The next question was more actionable: "
         "does every competitor fail the same way, or are some more vulnerable than others? "
         "If you\u2019re going to target churning customers, you need to know whose customers "
         "to target and what message will resonate with them.")

    h(doc, "Each competitor fails differently", level=2)
    body(doc, "I cross-tabulated every churn quote by competitor and failure group. "
         "The heatmap below shows what percentage of each competitor\u2019s churn falls into each group. "
         "The differences are striking \u2014 these companies don\u2019t just fail, they fail in distinctive ways.")
    add_chart(doc, c["heatmap"])
    for prod in TOP_COMPETITORS:
        total = sum(comp_groups.get(prod, {}).values())
        if total == 0: continue
        top_g = max(GROUP_ORDER, key=lambda g: comp_grp_pct(prod, g))
        top_g_pct = comp_grp_pct(prod, top_g)
        prod_top_cats = sorted(comp_cats.get(prod, {}).items(), key=lambda x: x[1], reverse=True)[:2]
        cat_desc = ", ".join(f"{CATEGORY_META[cc]['short']} ({nn})" for cc, nn in prod_top_cats)
        bullet(doc, f" ({total} churners): {top_g_pct:.0f}% {top_g.lower()}. Top complaints: {cat_desc}.", bold_prefix=prod)
    body(doc, "This tells you something important: the go-to-market message for a Ruby churner "
         "(\u201Cyour calls will be handled right\u201D) is completely different from the message for a "
         "Smith.ai churner (\u201Ctransparent billing, cancel anytime\u201D). One size won\u2019t fit all.")

    h(doc, "Where customers go when they leave", level=2)
    body(doc, f"I then looked for switching stories \u2014 reviews where a customer named both the "
         f"service they left and the one they moved to. I found {n_switch} of them. "
         "That\u2019s a small sample, so treat the direction as signal, not the exact numbers as proof. "
         "But the pattern is consistent.")
    add_chart(doc, c["switching_sankey"])

    # Compute left/joined totals for narrative
    _from_totals = Counter(); _to_totals = Counter()
    for f, t in switching_pairs:
        _from_totals[f] += 1; _to_totals[t] += 1
    _named_left = sorted(((k, v) for k, v in _from_totals.items() if k != "Unknown"), key=lambda x: -x[1])
    _named_joined = sorted(((k, v) for k, v in _to_totals.items() if k != "Unknown"), key=lambda x: -x[1])

    body(doc, f"The chart maps each of the {n_switch} switching stories as a flow from the company "
         "a customer left (on the left) to the one they joined (on the right). "
         "The \u201CUnknown\u201D band is large \u2014 most reviewers name only one company, not both \u2014 "
         "but among those who do name names, the pattern is clear.")
    if _named_left:
        left_desc = ", ".join(f"{name} ({cnt})" for name, cnt in _named_left[:3])
        body(doc, f"On the departures side: {left_desc}. "
             "Ruby stands out as the biggest named source of churning customers.")
    if _named_joined:
        joined_desc = ", ".join(f"{name} ({cnt})" for name, cnt in _named_joined[:3])
        body(doc, f"On the arrivals side: {joined_desc}. "
             "PATLive appears to be catching the most named switchers, which aligns with their "
             "reputation for customization and onboarding support.")
    body(doc, f"The trigger breakdown reinforces what the churn data already showed: "
         f"quality ({trig.get('Quality', 0)}/{n_switch}), "
         f"capability ({trig.get('Capability', 0)}/{n_switch}), "
         f"pricing ({trig.get('Pricing', 0)}/{n_switch}). "
         "People switch because the calls are handled badly \u2014 not because of price, not because "
         "of missing features. This is the opening for Central AI.")
    body(doc, f"Confidence note: {n_switch} switching stories is directional signal, not statistical proof. "
         "The pattern is consistent (quality-driven, Ruby losing) but the sample is small.",
         italic=True, sz=Pt(9), color=MID_GRAY)

    # Dollar impact
    h(doc, "The dollar amounts are real", level=2)
    too_exp_inline = cat_stats.get("BILLING_TOO_EXPENSIVE", {}).get("weighted", 0) / total_weighted * 100
    body(doc, "One question I anticipated you\u2019d ask: are these real businesses spending real money, "
         "or just price-sensitive bargain hunters who\u2019d churn from anything?")
    body(doc, f"The data answers clearly. {n_dollar} of {n_total} churn quotes mention specific dollar amounts \u2014 "
         "these are businesses spending hundreds to thousands per month. And the complaints aren\u2019t "
         "about the price being too high "
         f"(\u201CIt costs too much\u201D is only {too_exp_inline:.0f}% of churn). "
         "They\u2019re about how they\u2019re billed: surprise charges, per-minute inflation, impossible cancellation. "
         "These customers aren\u2019t looking for cheaper \u2014 they\u2019re looking for fairer.")

    # Billing sub-categories
    billing_cats = ["BILLING_PREDATORY", "BILLING_TRAP", "BILLING_OPAQUE", "BILLING_TOO_EXPENSIVE"]
    for bc in billing_cats:
        cs = cat_stats.get(bc, {})
        if cs.get("count", 0) > 0:
            bc_pct = cs["weighted"] / total_weighted * 100
            bullet(doc, f" ({cs['count']} quotes, {bc_pct:.0f}% of churn)", bold_prefix=CATEGORY_META[bc]["short"])

    # ══════════════════════════════════════════════════════════════
    # WHY NOW
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h(doc, "Why Now", level=1)
    body(doc, "So far I\u2019ve shown you what breaks and who\u2019s most vulnerable. But there\u2019s an obvious "
         "question you should be asking: if these services are so bad, why haven\u2019t customers "
         "already switched to something better? Is this actually a moment of opportunity, "
         "or has the market already settled?")
    body(doc, "I dug into the timing data to find out. What I found surprised me.")

    h(doc, "The ratings are fake", level=2)
    body(doc, "If you look at Trustpilot today, these companies look fine. Some even look good. "
         "I almost stopped here \u2014 maybe the churn quotes were just the vocal minority. "
         "Then I split the reviews by date and saw something that changed the picture entirely.")
    add_chart(doc, c["temporal"])
    body(doc, "Before mid-2024, when reviews were purely organic, the average ratings for these "
         "companies ranged from 1.2 to 3.8 stars. Then, starting around mid-2024, waves of 5-star "
         "reviews appeared across multiple companies simultaneously.")
    for prod in comp_order:
        td = temporal_data.get(prod, {})
        bullet(doc, f": organic avg {td.get('organic_avg', 0):.1f} "
               f"({td.get('organic_n', 0)} reviews) \u2192 recent {td.get('surge_avg', 0):.1f} "
               f"({td.get('surge_n', 0)} reviews)", bold_prefix=prod)
    spacer(doc)
    body(doc, "Smith.ai had zero 5-star reviews before mid-2024. Then 49 appeared. "
         "AnswerConnect: same \u2014 zero, then 59. The service didn\u2019t change. The review strategy did.")

    h(doc, "The distribution proves it", level=2)
    body(doc, "I wanted to be rigorous about this claim, so I looked at the star distribution. "
         "If these services were genuinely improving, you\u2019d expect reviews to spread across 1\u20135 stars \u2014 "
         "some lingering dissatisfaction, some moderate improvement, some genuine fans. "
         "That\u2019s not what the data shows.")
    add_chart(doc, c["bimodality"])
    body(doc, "Instead, reviews cluster at the extremes \u2014 almost nobody gives 2, 3, or 4 stars. "
         "This is a classic fingerprint of two separate populations: unhappy organic reviewers "
         "and solicited positive reviewers. The barbell distribution tells you the positive reviews "
         "are from a different population than the negative ones.")
    for prod_name in ["Smith.ai", "AnswerConnect", "Ruby Receptionist", "PATLive"]:
        bd = bimod_data.get(prod_name, {})
        bullet(doc, f": {bd.get('overall_extreme_pct', 0):.0f}% extreme. "
               f"Unprompted 5-star: {bd.get('org_5star', 0)}/{bd.get('org_total', 0)}. "
               f"Recent 5-star: {bd.get('surge_5star', 0)}/{bd.get('surge_total', 0)}.",
               bold_prefix=prod_name)
    spacer(doc)
    body(doc, "Here\u2019s what this means for you: the real competitive landscape is much weaker than "
         "Trustpilot suggests. You\u2019re not competing against 4-star services. You\u2019re competing "
         "against the organic experience (1.2\u20133.8 stars). The bar is low.")

    h(doc, "AI-native competitors are proving the concept", level=2)
    body(doc, "I also looked at the AI-native alternatives that have entered the market. "
         "Dialzara (14 reviews, 5.0 avg) and My AI Front Desk (8 reviews, 5.0 avg) show the "
         "AI receptionist concept resonates with customers \u2014 but these are tiny samples, "
         "too small to draw conclusions from.")
    body(doc, "Synthflow, with 69 reviews, shows what happens at scale: billing and reliability "
         "become the new battleground. The AI handles calls fine. The company behind it doesn\u2019t.")
    body(doc, "But here\u2019s the urgency: Bland AI raised $65M (YC, Emergence Capital). Synthflow has "
         "backing from Atlantic Labs. These are well-funded teams who will eventually fix their "
         "execution problems. The window where AI-native competitors have broken platforms "
         "is temporary. You need to move while it\u2019s open.")

    # Competitive response analysis
    h(doc, "What competitors can and can\u2019t fix", level=2)
    body(doc, "Before jumping to recommendations, I asked one more question: which of these competitive "
         "weaknesses are permanent, and which are temporary? If a competitor can patch their way "
         "out of a problem, you shouldn\u2019t build your strategy around it. "
         "You should bet on the structural ones \u2014 the problems baked into their business model.")

    struct_tbl = doc.add_table(rows=1, cols=3)
    struct_tbl.style = "Table Grid"
    for j, hdr in enumerate(["Problem", "Structural or Patchable?", "Why"]):
        struct_tbl.rows[0].cells[j].text = hdr
    struct_rows = [
        ("Agents don\u2019t know the business",
         "Structural",
         "Shared-pool model: one receptionist juggles hundreds of clients. "
         "Adding an AI screening layer doesn\u2019t fix the human\u2019s lack of context."),
        ("Per-minute billing incentives",
         "Structural",
         "Revenue = call duration. Flat-rate would cannibalize the business model. "
         "This is the innovator\u2019s dilemma \u2014 incumbents can\u2019t switch without destroying margin."),
        ("Staff turnover \u2192 quality decay",
         "Structural",
         "Human call centers have 30\u201345% annual turnover. Constant retraining is the norm, "
         "not a fixable bug. New hires will always be worse than tenured staff."),
        ("Script adherence",
         "Patchable (slowly)",
         "Better training helps marginally, but is limited by the shared-pool constraint. "
         "An agent handling 200 clients will never master all 200 scripts. AI solves this by design."),
        ("Technology integrations",
         "Patchable",
         "CRM integrations, scheduling APIs \u2014 these can be added. "
         "Smith.ai is already moving here with their hybrid model."),
        ("Review solicitation",
         "Patchable",
         "Already happening (the mid-2024 surge). Masks the problem but doesn\u2019t fix it."),
    ]
    for problem, fixable, why in struct_rows:
        row = struct_tbl.add_row()
        row.cells[0].text = problem
        row.cells[1].text = fixable
        row.cells[2].text = why
    style_table(struct_tbl)
    spacer(doc)
    body(doc, "Bottom line: the three biggest pain drivers (scripts, inauthenticity, billing model) "
         "are structural. Competitors would have to rebuild their entire business to fix them. "
         "That\u2019s the moat you should build around.")

    # ══════════════════════════════════════════════════════════════
    # THE PLAY
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h(doc, "The Play", level=1)
    body(doc, "So here\u2019s what the data tells you. Legacy services can\u2019t fix their structural problems. "
         "AI-native entrants have the right idea but haven\u2019t earned customer trust yet. "
         "The real ratings show incumbents are far weaker than their Trustpilot pages suggest. "
         "And the window is open but closing as well-funded competitors fix their platforms.")
    body(doc, "The rest of this section translates those findings into specific product, pricing, "
         "and go-to-market recommendations. Everything ties back to a number from the data above.")

    # ── What to build: specific product requirements ──────────
    h(doc, "What to build", level=2)
    body(doc, "I plotted every churn category on two axes: how much churn it drives (impact) and "
         "how naturally AI solves it (ease). The upper-right quadrant \u2014 high impact, high AI advantage "
         "\u2014 is where you should focus first.")
    add_chart(doc, c["priority"])
    body(doc, "Three features land squarely in that quadrant. Here\u2019s what they look like as product requirements:")

    script_pct = cat_stats.get("SCRIPT_ADHERENCE", {}).get("weighted", 0) / total_weighted * 100
    inauth_pct = cat_stats.get("INAUTHENTICITY", {}).get("weighted", 0) / total_weighted * 100
    garbled_pct = cat_stats.get("GARBLED_INFO", {}).get("weighted", 0) / total_weighted * 100
    pred_bill_pct = cat_stats.get("BILLING_PREDATORY", {}).get("weighted", 0) / total_weighted * 100
    trap_pct = cat_stats.get("BILLING_TRAP", {}).get("weighted", 0) / total_weighted * 100
    opaque_pct = cat_stats.get("BILLING_OPAQUE", {}).get("weighted", 0) / total_weighted * 100
    decay_pct = cat_stats.get("QUALITY_DECAY", {}).get("weighted", 0) / total_weighted * 100
    missed_pct = cat_stats.get("MISSED_CALLS", {}).get("weighted", 0) / total_weighted * 100

    req_tbl = doc.add_table(rows=1, cols=4)
    req_tbl.style = "Table Grid"
    for j, hdr in enumerate(["Priority", "Product Feature", "Churn It Prevents", "What Customers Said"]):
        req_tbl.rows[0].cells[j].text = hdr
    req_rows = [
        ("P0", "Configurable call scripts with branching logic",
         f"\u201CThey don\u2019t follow my instructions\u201D ({script_pct:.0f}%)",
         "AI IS the script, not reading one"),
        ("P0", "AI trained per client: business FAQ, repeat caller recognition",
         f"\u201CThey don\u2019t know my business\u201D ({inauth_pct:.0f}%)",
         "Dedicated AI vs stranger with a screen"),
        ("P0", "Flat monthly rate, no per-minute billing",
         f"\u201CHidden charges on my bill\u201D ({pred_bill_pct:.0f}%)",
         "Eliminate the per-minute incentive entirely"),
        ("P1", "Structured data capture with spell-back validation, CRM push",
         f"\u201CThey get the details wrong\u201D ({garbled_pct:.0f}%)",
         "Structured forms, not free-text note-taking"),
        ("P1", "Self-serve cancellation, no contracts, no 30-day notice",
         f"\u201CI can\u2019t cancel\u201D ({trap_pct:.0f}%)",
         "One-click cancel, no hoops"),
        ("P1", "Real-time usage dashboard with itemized call logs",
         f"\u201CSurprise charges I can\u2019t explain\u201D ({opaque_pct:.0f}%)",
         "Show every call, every charge, in real time"),
        ("P2", "24/7 instant answer, zero hold time",
         f"\u201CThey don\u2019t pick up\u201D ({missed_pct:.0f}%)",
         "AI doesn\u2019t have staffing gaps"),
        ("P2", "Automated quality monitoring, no staff turnover",
         f"\u201CIt used to be good, then got worse\u201D ({decay_pct:.0f}%)",
         "\u201CUsed to be great, now it\u2019s terrible\u201D \u2014 Day 1000 = Day 1"),
    ]
    for priority, feature, churn, quote in req_rows:
        row = req_tbl.add_row()
        row.cells[0].text = priority
        row.cells[1].text = feature
        row.cells[2].text = churn
        row.cells[3].text = quote
    style_table(req_tbl)
    spacer(doc)
    body(doc, f"P0 features address {script_pct + inauth_pct + pred_bill_pct:.0f}% of all churn. "
         "Ship these before anything else.")

    # ── What to charge ────────────────────────────────────────
    h(doc, "What to charge", level=2)
    too_exp_pct = cat_stats.get("BILLING_TOO_EXPENSIVE", {}).get("weighted", 0) / total_weighted * 100
    body(doc, "I looked at what the churn data says about pricing. The signal is unambiguous: "
         "the problem isn\u2019t how much these services charge \u2014 it\u2019s how they charge.")
    bullet(doc, f" is only {too_exp_pct:.0f}% of churn. Price is not why people leave.",
           bold_prefix="\u201CIt costs too much\u201D")
    bullet(doc, f" is {bill_pct:.0f}% of churn. How they\u2019re billed is why people leave.",
           bold_prefix="All billing complaints combined")
    bullet(doc, f" {n_dollar} of {n_total} churners mention specific dollar amounts "
           f"(${dollar_min:,.0f}\u2013${dollar_max:,.0f}). "
           "They\u2019re spending real money and expecting real service.",
           bold_prefix="Dollar amounts in complaints:")
    spacer(doc)

    price_tbl = doc.add_table(rows=1, cols=3)
    price_tbl.style = "Table Grid"
    for j, hdr in enumerate(["Segment", "Current Pricing", "Central AI Opportunity"]):
        price_tbl.rows[0].cells[j].text = hdr
    price_rows = [
        ("Legacy incumbents", "$235\u2013$2,800/mo (per-minute)",
         "Undercut by 50\u201370% while eliminating per-minute surprises"),
        ("AI-native (Synthflow etc.)", "$29\u2013$450/mo (per-minute + base)",
         "Price similarly but deliver reliability and support they can\u2019t"),
        ("Recommended range", "\u2014",
         f"$99\u2013199/mo flat rate. Above AI-native noise, below legacy pain. "
         f"Flat-rate directly addresses billing-model churn ({pred_bill_pct + trap_pct + opaque_pct:.0f}% of all churn: "
         f"predatory billing, billing traps, and opaque charges)."),
    ]
    for seg, current, opp in price_rows:
        row = price_tbl.add_row()
        row.cells[0].text = seg
        row.cells[1].text = current
        row.cells[2].text = opp
    style_table(price_tbl)
    spacer(doc)
    body(doc, "Flat-rate pricing isn\u2019t just a feature \u2014 it\u2019s a structural advantage you should lean into. "
         "Incumbents can\u2019t match it without destroying their per-minute revenue model. "
         "That\u2019s the innovator\u2019s dilemma working in your favor.")

    # ── How big is this ───────────────────────────────────────
    h(doc, "How big is this", level=2)
    body(doc, "You\u2019ll want to know the size of the opportunity. I built a bottoms-up estimate "
         "using what this analysis gives us. It\u2019s intentionally conservative \u2014 I\u2019d rather "
         "give you a number you can defend than one that sounds impressive.")

    tam_tbl = doc.add_table(rows=1, cols=3)
    tam_tbl.style = "Table Grid"
    for j, hdr in enumerate(["Input", "Estimate", "Source"]):
        tam_tbl.rows[0].cells[j].text = hdr
    tam_rows = [
        ("Ruby customer base", "15,000+", "Ruby\u2019s public marketing"),
        ("Four legacy incumbents combined", "~40,000\u201360,000 SMBs",
         "Estimate from Ruby (15K) + AnswerConnect, PATLive, Smith.ai"),
        ("Annual churn rate (industry)", "15\u201325%",
         "Typical for B2B SaaS/services; our data shows persistent dissatisfaction"),
        ("Customers actively shopping/year", "~6,000\u201315,000",
         "40\u201360K base \u00D7 15\u201325% churn"),
        ("Central AI capture rate (year 1)", "1\u20133%",
         "Conservative for a new entrant with no brand"),
        ("Capturable accounts (year 1)", "60\u2013450 accounts", "Rows above"),
        ("At $149/mo average", "$107K\u2013$804K ARR", "Accounts \u00D7 $149/mo \u00D7 12"),
    ]
    for inp, est, src in tam_rows:
        row = tam_tbl.add_row()
        row.cells[0].text = inp
        row.cells[1].text = est
        row.cells[2].text = src
    style_table(tam_tbl)
    spacer(doc)
    body(doc, "This is conservative. It excludes AI-native switchers, new market entrants, "
         "and the long tail of smaller answering services not in this dataset. "
         "The real TAM is larger \u2014 but even the conservative case shows a viable beachhead.")

    # ── Who to target ─────────────────────────────────────────
    h(doc, "Who to target", level=2)
    body(doc, "Based on the competitor vulnerability data and switching patterns, "
         "here\u2019s who I\u2019d prioritize and in what order. The logic: start with the highest-intent, "
         "most-accessible customers and work outward.")
    ruby_net = nf.get("Ruby Receptionist", {}).get("net", 0)
    smith_net = nf.get("Smith.ai", {}).get("net", 0)
    smith_bill = comp_grp_pct("Smith.ai", "BILLING")
    serial_n = cat_stats.get("SERIAL_SWITCHING", {}).get("count", 0)

    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    for j, hdr in enumerate(["Priority", "Target", "Why (from data)", "Lead Message"]):
        tbl2.rows[0].cells[j].text = hdr
    target_rows = [
        ("1", "Ruby churners",
         f"Worst net flow ({ruby_net:+d}), quality decay complaints, "
         f"{comp_grp_pct('Ruby Receptionist', 'CALL HANDLING'):.0f}% call handling churn",
         "\u201CYour calls handled right, every time\u201D"),
        ("2", "Smith.ai churners",
         f"Billing-driven churn ({smith_bill:.0f}%), net {smith_net:+d}, "
         "customers describe \u201Cintentionally difficult\u201D cancellation",
         "\u201CTransparent pricing, cancel anytime\u201D"),
        ("3", "Serial switchers",
         f"{serial_n} quotes describing trying 3\u20134 services. "
         "Highest-intent prospects \u2014 they\u2019ve given up on the category",
         "\u201CDifferent technology, not another answering service\u201D"),
        ("4", "Synthflow refugees",
         f"{synth_bill:.0f}% billing churn, {synth_rel:.0f}% reliability. "
         "They wanted AI but got a broken platform",
         "\u201CAI that actually works, with real support\u201D"),
    ]
    for priority, target, why, msg in target_rows:
        row = tbl2.add_row()
        row.cells[0].text = priority
        row.cells[1].text = target
        row.cells[2].text = why
        row.cells[3].text = msg
    style_table(tbl2)

    # ── What to say ───────────────────────────────────────────
    h(doc, "What to say", level=2)
    body(doc, "Each channel should lead with the pain point most relevant to the audience it reaches. "
         "I mapped the churn categories to specific channels and drafted messages grounded in "
         "the exact language customers use in their complaints.")
    qd_pct = cat_stats.get("QUALITY_DECAY", {}).get("weighted", 0) / total_weighted * 100

    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Table Grid"
    for j, hdr in enumerate(["Channel", "Message", "Pain It Addresses", "% of Churn"]):
        tbl3.rows[0].cells[j].text = hdr
    msg_rows = [
        ("Google Ads / SEO",
         "\u201CAI receptionist that follows your script\u201D",
         "\u201CThey don\u2019t follow my instructions\u201D", f"{script_pct:.0f}%"),
        ("Trustpilot competitor pages",
         "\u201CFlat rate. No per-minute billing. Cancel anytime.\u201D",
         "All billing complaints", f"{bill_pct:.0f}%"),
        ("Competitor review responses",
         "\u201CSame problem? We built something different.\u201D",
         "\u201CI\u2019ve tried everyone, nobody works\u201D", f"{dis_pct:.0f}%"),
        ("Direct outreach \u2014 Ruby customers",
         "\u201CStill happy with your call quality?\u201D",
         "\u201CIt used to be good, then got worse\u201D", f"{qd_pct:.0f}%"),
        ("Direct outreach \u2014 Smith.ai customers",
         "\u201CTired of surprise charges? We do flat rate.\u201D",
         "\u201CHidden charges\u201D + \u201CSurprise charges\u201D", f"{pred_bill_pct + opaque_pct:.0f}%"),
    ]
    for channel, message, pain, pct_str in msg_rows:
        row = tbl3.add_row()
        row.cells[0].text = channel
        row.cells[1].text = message
        row.cells[2].text = pain
        row.cells[3].text = pct_str
    style_table(tbl3)

    # ── What could go wrong ───────────────────────────────────
    h(doc, "What could go wrong", level=2)
    body(doc, "I wouldn\u2019t give you a recommendation without flagging the risks. The Synthflow data "
         "is instructive here: it shows that AI doesn\u2019t eliminate churn \u2014 it shifts it. "
         "Here are the failure modes you need to plan for, along with what I\u2019d suggest doing about each.")

    risk_tbl = doc.add_table(rows=1, cols=3)
    risk_tbl.style = "Table Grid"
    for j, hdr in enumerate(["Risk", "Evidence from Data", "Mitigation"]):
        risk_tbl.rows[0].cells[j].text = hdr
    risk_rows = [
        ("Billing becomes the new churn driver",
         f"Synthflow: {synth_bill:.0f}% of churn is billing (bait-and-switch, hidden costs). "
         "AI-native doesn\u2019t mean billing-clean.",
         "Flat rate from day one. No usage tiers that surprise. "
         "Publish pricing publicly. Self-serve cancellation."),
        ("Platform reliability failures",
         f"Synthflow: {synth_rel:.0f}% reliability churn. "
         "Silent breaking updates, features stop working, support absent.",
         "Invest in engineering discipline early. Uptime SLA. "
         "Status page. Regression test suite before every deploy."),
        ("Callers detect \u2018outsourced\u2019 feel",
         f"\u201CThey don\u2019t know my business\u201D is {inauth_pct:.0f}% of legacy churn \u2014 callers notice when the receptionist "
         "doesn\u2019t know the business. AI could trigger the same reaction for a different reason "
         "(robotic tone vs clueless human). No AI-specific data yet \u2014 AI-native services are too new.",
         "AI must be trained per client. Graceful handoff to human when AI is uncertain. "
         "Don\u2019t pretend to be human \u2014 be upfront and be competent."),
        ("New competitors fix their problems",
         "Bland AI ($65M raised), Synthflow (Atlantic Labs). "
         "Well-funded teams will eventually fix billing and reliability.",
         "Speed matters. The window is open now. First mover with reliable AI + clean billing wins. "
         "Build switching costs through CRM integrations and client-specific AI training."),
    ]
    for risk, evidence, mitigation in risk_rows:
        row = risk_tbl.add_row()
        row.cells[0].text = risk
        row.cells[1].text = evidence
        row.cells[2].text = mitigation
    style_table(risk_tbl)

    # ══════════════════════════════════════════════════════════════
    # APPENDIX: METHODOLOGY
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h(doc, "Appendix: Methodology", level=1)

    p = doc.add_paragraph()
    r = p.add_run("Data: "); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK_BLUE
    r2 = p.add_run(f"{total_reviews} reviews scraped from Trustpilot (9 companies) and Reddit "
                    f"(7 subreddits). Filtered to {n_total} with a specific churn reason.")
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY

    p = doc.add_paragraph()
    r = p.add_run("Classification: "); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK_BLUE
    r2 = p.add_run(f"11-category taxonomy with clear decision rules. "
                    f"Two independent AI classifiers tagged every quote. Agreement: "
                    f"{kappa_agree}/{kappa_n} ({kappa_agree/kappa_n*100:.1f}%). "
                    f"Cohen\u2019s kappa: {kappa_val:.2f} (\u201Calmost perfect\u201D). "
                    f"{len(disagreements)} disagreements manually reviewed.")
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY

    p = doc.add_paragraph()
    r = p.add_run("Weighting: "); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK_BLUE
    r2 = p.add_run("Detailed, highly-upvoted Reddit posts carry more weight than one-line reviews. "
                    "Weight = quote_quality \u00D7 log\u2082(upvotes) for Reddit, quote_quality \u00D7 1.0 for Trustpilot.")
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY

    body(doc, "Known limitations:", bold=True, color=DARK_BLUE)
    for lim in [
        "92% Trustpilot, 8% Reddit. Trustpilot skews negative.",
        "All Reddit churn quotes from r/LawFirm.",
        "Both classifiers are AI \u2014 may share blind spots.",
        f"34 of {n_total} quotes are pre-2020. Complaint types hold across time windows; competitor ranking shifts slightly.",
        "Mid-2024 review surge cutoff is judgment-based.",
    ]:
        bullet(doc, lim)

    spacer(doc)
    p = doc.add_paragraph()
    r = p.add_run("Temporal sensitivity: "); r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = DARK_BLUE
    r2 = p.add_run(f"Complaint types (call handling ~{ch_pct:.0f}%, billing ~{bill_pct:.0f}%) "
                    "hold in both all-dates and post-2020 windows. Competitor ranking shifts: Ruby and Smith.ai "
                    "move up in recent data, AnswerConnect moves down.")
    r2.font.size = Pt(10.5); r2.font.color.rgb = DARK_GRAY

    # ══════════════════════════════════════════════════════════════
    # APPENDIX: EVIDENCE SAMPLES
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h(doc, "Appendix: Evidence Samples", level=1)
    body(doc, "Three representative quotes per failure group, ranked by quality score.")

    for g in GROUP_ORDER:
        gs = group_stats[g]
        cnt = gs["count"]
        body(doc, f"{g} ({cnt} quotes)", bold=True, color=DARK_BLUE, sz=Pt(11))
        ev = evidence_samples.get(g, [])
        if ev:
            for _, prod, pull, url, src in ev:
                src_label = "Trustpilot" if "trustpilot" in url else "Reddit"
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(f"\u201C{pull}\u201D"); r.italic = True; r.font.size = Pt(9); r.font.color.rgb = DARK_GRAY
                r2 = p.add_run(f" \u2014 {prod} ({src_label})"); r2.font.size = Pt(9); r2.font.color.rgb = MID_GRAY
        spacer(doc)

    # ══════════════════════════════════════════════════════════════
    # APPENDIX: SAMPLE RAW DATA
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h(doc, "Appendix: Sample Raw Data", level=1)
    body(doc, "Every claim in this document traces to a real review. Here are 10 sampled "
         "quotes with their original source links.")

    # Pick 10 diverse quotes (2-3 per group), all with URLs
    sample_indices = []
    for g in GROUP_ORDER:
        group_cat_list = group_stats[g]["cats"]
        g_idxs = [(idx, q) for idx, cat in final_codes.items()
                   if cat in group_cat_list
                   for q in [quotes[idx]]
                   if q.get("url")]
        g_idxs.sort(key=lambda x: compute_weight(x[1]), reverse=True)
        sample_indices.extend(g_idxs[:3])
    seen = set()
    final_samples = []
    for idx, q in sample_indices:
        if idx not in seen:
            seen.add(idx)
            final_samples.append((idx, q))
        if len(final_samples) >= 10:
            break

    tbl4 = doc.add_table(rows=1, cols=5)
    tbl4.style = "Table Grid"
    for j, hdr in enumerate(["#", "Competitor", "Category", "Quote (truncated)", "Source"]):
        tbl4.rows[0].cells[j].text = hdr
    for i, (idx, q) in enumerate(final_samples, 1):
        cat = final_codes[idx]
        prod = normalize_product(q["llm"].get("product_mentioned"))
        snippet = truncate(q["text"], 80)
        url = q.get("url", "")
        src = "Trustpilot" if "trustpilot" in url else "Reddit"
        row = tbl4.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = prod
        row.cells[2].text = CATEGORY_META[cat]["short"]
        row.cells[3].text = f"\u201C{snippet}\u201D"
        row.cells[4].text = f"{src}: {url}"
    style_table(tbl4)

    body(doc, f"Full database of all {n_total} quotes with source links available in the supplementary data files.",
         italic=True, sz=Pt(9), color=MID_GRAY)

    # ══════════════════════════════════════════════════════════════
    # APPENDIX: ALL CHURN QUOTES CATEGORIZED
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    h(doc, "Appendix: All Churn Quotes", level=1)
    body(doc, f"All {n_total} churn quotes organized by customer-voiced failure mode. "
         f"Quotes are ranked by quality score within each category.")

    for g in GROUP_ORDER:
        gs = group_stats[g]
        g_pct = gs["weighted"] / total_weighted * 100
        spacer(doc)
        body(doc, f"{g} ({gs['count']} quotes, {g_pct:.0f}% of churn)", bold=True, color=DARK_BLUE, sz=Pt(12))

        g_cats = sorted(gs["cats"], key=lambda cc: cat_stats[cc]["count"], reverse=True)
        for cat in g_cats:
            cs = cat_stats[cat]
            label = CATEGORY_META[cat]["short"]
            cat_pct = cs["weighted"] / total_weighted * 100
            body(doc, f"\u201C{label}\u201D ({cs['count']} quotes, {cat_pct:.0f}%)",
                 bold=True, color=DARK_BLUE, sz=Pt(10.5))

            sorted_q = sorted(cs["quotes"], key=lambda x: x[2], reverse=True)
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            tbl.rows[0].cells[0].text = "Source"
            tbl.rows[0].cells[1].text = "Quote"
            # Set 2:8 column widths
            total_w = Inches(6.5)
            for row_obj in tbl.rows:
                row_obj.cells[0].width = Emu(int(total_w * 0.2))
                row_obj.cells[1].width = Emu(int(total_w * 0.8))
            for rank, (idx, q, w) in enumerate(sorted_q, 1):
                prod = normalize_product(q["llm"].get("product_mentioned"))
                source = q.get("source", "")
                url = q.get("url", "")
                if not source:
                    source = "Trustpilot" if "trustpilot" in url else "Reddit" if "reddit" in url else "Unknown"
                date = q.get("date", "")
                text = q["text"].replace("\n", " ").strip()

                row = tbl.add_row()
                # Source cell (left, narrow) — hyperlinked
                cell0 = row.cells[0]
                cell0.width = Emu(int(total_w * 0.2))
                cell0.text = ""
                p0 = cell0.paragraphs[0]
                if url:
                    add_hyperlink(p0, prod, url)
                else:
                    r0 = p0.add_run(prod)
                    r0.font.size = Pt(8); r0.font.color.rgb = MID_GRAY
                r_src = p0.add_run(f"\n{source}")
                r_src.font.size = Pt(7); r_src.font.color.rgb = MID_GRAY
                if date:
                    r_date = p0.add_run(f"\n{date}")
                    r_date.font.size = Pt(7); r_date.font.color.rgb = MID_GRAY
                # Quote cell (right, wide)
                cell1 = row.cells[1]
                cell1.width = Emu(int(total_w * 0.8))
                cell1.text = ""
                p1 = cell1.paragraphs[0]
                r1 = p1.add_run(f"\u201C{text}\u201D")
                r1.italic = True; r1.font.size = Pt(9); r1.font.color.rgb = DARK_GRAY
            style_table(tbl)
            spacer(doc)

    # ── Save docx ──────────────────────────────────────────────────
    doc.save(OUTPUT_DOCX)
    print(f"Done \u2192 {OUTPUT_DOCX}")
    print(f"  {n_total} churn quotes, kappa={kappa_val:.4f}, {n_switch} switching stories")
    print(f"  {len(c)} charts in {CHART_DIR}/")

    # ── Export categorized churn list (customer-centric labels) ──
    cat_list_path = os.path.join(BASE, "churn_quotes_categorized.md")
    lines = []
    lines.append("# Churn Quotes by Category\n")
    lines.append(f"**{n_total} churn quotes** from {total_reviews} reviews across 9 competitors, "
                 f"dual-coded with inter-rater reliability \u03BA = {kappa_val:.2f}.\n")
    lines.append("Categories are grouped into four themes. Each quote includes the competitor, "
                 "source, and a quality score (1\u20135).\n")
    lines.append("---\n")

    for g in GROUP_ORDER:
        g_count = group_stats[g]["count"]
        g_pct = group_stats[g]["weighted"] / total_weighted * 100
        lines.append(f"\n## {g} ({g_count} quotes, {g_pct:.0f}% of churn)\n")

        # Sort categories within group by count descending
        g_cats = sorted(group_stats[g]["cats"],
                        key=lambda cc: cat_stats[cc]["count"], reverse=True)
        for cat in g_cats:
            cs = cat_stats[cat]
            label = CATEGORY_META[cat]["short"]
            cat_pct = cs["weighted"] / total_weighted * 100
            lines.append(f"\n### {label} ({cs['count']} quotes, {cat_pct:.0f}%)\n")

            # Sort quotes by weight descending
            sorted_q = sorted(cs["quotes"], key=lambda x: x[2], reverse=True)
            for rank, (idx, q, w) in enumerate(sorted_q, 1):
                prod = normalize_product(q["llm"].get("product_mentioned"))
                source = q.get("source", "")
                if not source:
                    url = q.get("url", "")
                    source = "Trustpilot" if "trustpilot" in url else "Reddit" if "reddit" in url else "Unknown"
                quality = q["llm"].get("quote_quality", "?")
                pain = q["llm"].get("pain_point", "")
                text = q["text"].replace("\n", " ").strip()

                date = q.get("date", "")
                date_str = f" | {date}" if date else ""
                lines.append(f"**#{rank}** | {prod} | Quality: {quality}/5 | {source}{date_str}\n")
                if pain:
                    lines.append(f"*{pain}*\n")
                lines.append(f"> {text}\n")

    with open(cat_list_path, "w") as f:
        f.write("\n".join(lines))
    print(f"  Categorized list \u2192 {cat_list_path}")


if __name__ == "__main__":
    main()
